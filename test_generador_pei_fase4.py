import contextlib
import os
import re
import shutil
import sys
import tempfile
import types
import unittest
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path
from unittest.mock import patch

import generador_pei_fase4
from generador_pei_fase4 import (
    PayloadValidationError,
    PEIHandler,
    construir_fichas_seleccionadas,
    construir_mapeo_codigos_display,
    extraer_anios_periodo,
    indexar_fichas,
    indexar_metas,
    limpiar_parrafos_body_vacios_finales,
    mapear_columnas_anio,
    ajustar_columnas_anio,
    asegurar_salto_de_pagina_antes,
    eliminar_parrafos_con_texto,
    extraer_nota_ficha,
    formatear_oracion_anexo_b1,
    normalizar_bordes_tabla,
    reemplazar_placeholder_docx,
    remapear_codigo_aei,
    remapear_codigos_visibles_docx,
    resolver_anios_columnas,
    resolver_linea_base_ficha,
    sanitizar_nombre_municipio,
    seleccionar_ficha_base,
    validar_payload,
)


MATRIZ = {
    'oei': [
        {
            'codigo': 'OEI.01',
            'indicadores': [{}],
            'aei': [
                {'codigo': 'AEI.01.01', 'indicadores': [{}]},
                {'codigo': 'AEI.01.02', 'indicadores': [{}]},
            ],
        },
        {
            'codigo': 'OEI.02',
            'indicadores': [{}, {}],
            'aei': [{'codigo': 'AEI.02.01', 'indicadores': [{}, {}, {}, {}, {}]}],
        },
    ]
}


def payload(oei, aei, priorities_oei, priorities_aei, indicators_oei=None, indicators_aei=None):
    return {
        'periodo_pei': '2026-2030',
        'ordenanza_pdc': '002-2024-MDL',
        'selecciones': {
            'oei': [f'oei-{code}' for code in oei],
            'aei': [f'aei-{code}' for code in aei],
            'indicadoresOEI': indicators_oei or [f'ind-oei-{code}-0' for code in oei],
            'indicadoresAEI': indicators_aei or [f'ind-aei-{code}-0' for code in aei],
        },
        'prioridades': {
            'oei': priorities_oei,
            'aei': priorities_aei,
        },
    }


class ValidarPayloadTests(unittest.TestCase):
    def test_rejects_empty_period(self):
        data = payload(
            ['OEI.01'], ['AEI.01.01'], ['OEI.01'], {'OEI.01': ['AEI.01.01']}
        )
        data['periodo_pei'] = ''

        with self.assertRaises(PayloadValidationError) as context:
            validar_payload(data, MATRIZ)

        self.assertIn('required', {error['code'] for error in context.exception.errors})

    def test_rejects_invalid_period_format(self):
        data = payload(
            ['OEI.01'], ['AEI.01.01'], ['OEI.01'], {'OEI.01': ['AEI.01.01']}
        )
        data['periodo_pei'] = '2027/2032'

        with self.assertRaises(PayloadValidationError) as context:
            validar_payload(data, MATRIZ)

        self.assertIn('invalid_period', {error['code'] for error in context.exception.errors})

    def test_rejects_ordinance_placeholder(self):
        data = payload(
            ['OEI.01'], ['AEI.01.01'], ['OEI.01'], {'OEI.01': ['AEI.01.01']}
        )
        data['ordenanza_pdc'] = '«Ordenanza_PDC»'

        with self.assertRaises(PayloadValidationError) as context:
            validar_payload(data, MATRIZ)

        self.assertIn('placeholder', {error['code'] for error in context.exception.errors})
    def test_preserves_priority_order(self):
        result = validar_payload(
            payload(
                ['OEI.01', 'OEI.02'],
                ['AEI.01.01', 'AEI.01.02', 'AEI.02.01'],
                ['OEI.02', 'OEI.01'],
                {
                    'OEI.01': ['AEI.01.02', 'AEI.01.01'],
                    'OEI.02': ['AEI.02.01'],
                },
            ),
            MATRIZ,
        )

        self.assertEqual(result['oei'], ['OEI.02', 'OEI.01'])
        self.assertEqual(result['aei']['OEI.01'], ['AEI.01.02', 'AEI.01.01'])

    def test_rejects_child_without_selected_parent(self):
        with self.assertRaises(PayloadValidationError) as context:
            validar_payload(
                payload(
                    ['OEI.01'],
                    ['AEI.02.01'],
                    ['OEI.01'],
                    {'OEI.01': ['AEI.02.01']},
                    indicators_aei=['ind-aei-AEI.02.01-0'],
                ),
                MATRIZ,
            )

        codes = {error['code'] for error in context.exception.errors}
        self.assertIn('missing_parent', codes)

    def test_rejects_invalid_priority_relation(self):
        with self.assertRaises(PayloadValidationError) as context:
            validar_payload(
                payload(
                    ['OEI.01'],
                    ['AEI.01.01'],
                    ['OEI.01'],
                    {'OEI.01': ['AEI.02.01']},
                ),
                MATRIZ,
            )

        codes = {error['code'] for error in context.exception.errors}
        self.assertIn('invalid_relation', codes)

    def test_accepts_multiple_indicator_ordinals(self):
        result = validar_payload(
            payload(
                ['OEI.02'],
                ['AEI.02.01'],
                ['OEI.02'],
                {'OEI.02': ['AEI.02.01']},
                indicators_oei=['ind-oei-OEI.02-0', 'ind-oei-OEI.02-1'],
                indicators_aei=[f'ind-aei-AEI.02.01-{index}' for index in range(5)],
            ),
            MATRIZ,
        )

        self.assertEqual(result['oei'], ['OEI.02'])
        self.assertEqual(result['aei']['OEI.02'], ['AEI.02.01'])

    def test_rejects_reversed_period(self):
        data = payload(
            ['OEI.01'],
            ['AEI.01.01'],
            ['OEI.01'],
            {'OEI.01': ['AEI.01.01']},
        )
        data['periodo_pei'] = '2030-2024'

        with self.assertRaises(PayloadValidationError) as context:
            validar_payload(data, MATRIZ)

        self.assertIn('invalid_period', {error['code'] for error in context.exception.errors})

    def test_rejects_invalid_codigo_ue(self):
        data = payload(
            ['OEI.01'],
            ['AEI.01.01'],
            ['OEI.01'],
            {'OEI.01': ['AEI.01.01']},
        )
        data['codigo_ue'] = 'UE-001'

        with self.assertRaises(PayloadValidationError) as context:
            validar_payload(data, MATRIZ)

        self.assertIn('invalid_format', {error['code'] for error in context.exception.errors})

    def test_rejects_non_numeric_meta_value(self):
        data = payload(
            ['OEI.01'],
            ['AEI.01.01'],
            ['OEI.01'],
            {'OEI.01': ['AEI.01.01']},
        )
        data['metas'] = {'ind-oei-OEI.01-0': {'meta_2026': []}}

        with self.assertRaises(PayloadValidationError) as context:
            validar_payload(data, MATRIZ)

        self.assertIn('invalid_type', {error['code'] for error in context.exception.errors})

    def test_rejects_negative_values_in_baseline_and_annual_targets(self):
        for field in ('valor_base', 'meta_2026'):
            with self.subTest(field=field):
                data = payload(
                    ['OEI.01'],
                    ['AEI.01.01'],
                    ['OEI.01'],
                    {'OEI.01': ['AEI.01.01']},
                )
                data['metas'] = {
                    'ind-oei-OEI.01-0': {
                        'año_base': 2024,
                        field: -1,
                    }
                }

                with self.assertRaises(PayloadValidationError) as context:
                    validar_payload(data, MATRIZ)

                errors = context.exception.errors
                self.assertIn('negative_value', {error['code'] for error in errors})
                self.assertIn(f'metas.ind-oei-OEI.01-0.{field}', {error['field'] for error in errors})

    def test_accepts_zero_positive_values_and_period_years(self):
        data = payload(
            ['OEI.01'],
            ['AEI.01.01'],
            ['OEI.01'],
            {'OEI.01': ['AEI.01.01']},
        )
        data['periodo_pei'] = '2027-2032'
        data['metas'] = {
            'ind-oei-OEI.01-0': {
                'año_base': 2031,
                'valor_base': 0,
                'meta_2027': 0,
                'meta_2032': 12.5,
            }
        }

        self.assertEqual(validar_payload(data, MATRIZ), {
            'oei': ['OEI.01'],
            'aei': {'OEI.01': ['AEI.01.01']},
        })


class DisplayCodeMappingTests(unittest.TestCase):
    def test_builds_sequential_display_codes_from_priority_order(self):
        ordered = ['OEI.01', 'OEI.04', 'OEI.10']

        self.assertEqual(construir_mapeo_codigos_display(ordered), {
            'OEI.01': 'OEI.01',
            'OEI.04': 'OEI.02',
            'OEI.10': 'OEI.03',
        })
        self.assertEqual(ordered, ['OEI.01', 'OEI.04', 'OEI.10'])

    def test_remaps_first_selected_aei_to_visible_one(self):
        mapping = construir_mapeo_codigos_display(
            ['OEI.01'],
            {'OEI.01': ['AEI.01.05']},
        )

        self.assertEqual(remapear_codigo_aei('AEI.01.05', mapping), 'AEI.01.01')

    def test_remaps_selected_aei_in_priority_order(self):
        aei_priority = ['AEI.01.05', 'AEI.01.07']
        mapping = construir_mapeo_codigos_display(
            ['OEI.01'],
            {'OEI.01': aei_priority},
        )

        self.assertEqual(remapear_codigo_aei('AEI.01.05', mapping), 'AEI.01.01')
        self.assertEqual(remapear_codigo_aei('AEI.01.07', mapping), 'AEI.01.02')
        self.assertEqual(aei_priority, ['AEI.01.05', 'AEI.01.07'])

    def test_remaps_aei_parent_and_local_priority(self):
        mapping = construir_mapeo_codigos_display(
            ['OEI.01', 'OEI.04', 'OEI.10'],
            {
                'OEI.01': ['AEI.01.01'],
                'OEI.04': ['AEI.04.02'],
                'OEI.10': ['AEI.10.01'],
            },
        )

        self.assertEqual(remapear_codigo_aei('AEI.04.02', mapping), 'AEI.02.01')
        self.assertEqual(remapear_codigo_aei('AEI.10.01', mapping), 'AEI.03.01')
        self.assertEqual(remapear_codigo_aei('AEI.02.01', mapping), 'AEI.02.01')

    def test_docx_code_remap_keeps_run_styles_and_internal_codes(self):
        from docx import Document

        document = Document()
        paragraph = document.add_paragraph()
        prefix = paragraph.add_run('Acciones del OEI.')
        prefix.bold = True
        suffix = paragraph.add_run('04 y AEI.04.02')
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = 'Ficha OEI.10 / AEI.10.01'
        internal_ids = document.add_paragraph('ind-oei-OEI.04-0 / oei-OEI.04')
        mapping = construir_mapeo_codigos_display(
            ['OEI.01', 'OEI.04', 'OEI.10'],
            {
                'OEI.01': ['AEI.01.01'],
                'OEI.04': ['AEI.04.02'],
                'OEI.10': ['AEI.10.01'],
            },
        )

        remapear_codigos_visibles_docx(document, mapping)

        self.assertEqual(paragraph.text, 'Acciones del OEI.02 y AEI.02.01')
        self.assertTrue(prefix.bold)
        self.assertEqual(suffix.text, ' y AEI.02.01')
        self.assertEqual(table.cell(0, 0).text, 'Ficha OEI.03 / AEI.03.01')
        self.assertEqual(internal_ids.text, 'ind-oei-OEI.04-0 / oei-OEI.04')


class DataCorrelationTests(unittest.TestCase):
    def test_metas_are_indexed_by_indicator_ordinal(self):
        metas = indexar_metas({
            'ind-aei-AEI.02.01-0': {'meta_2026': 10},
            'ind-aei-AEI.02.01-1': {'meta_2026': 20},
        })

        self.assertEqual(metas[('AEI.02.01', 0)]['meta_2026'], 10)
        self.assertEqual(metas[('AEI.02.01', 1)]['meta_2026'], 20)

    def test_base_year_aliases_are_normalized_to_anio_base(self):
        normalized = []
        for alias in ('año_base', 'anio_base'):
            meta = indexar_metas({
                'ind-oei-OEI.01-0': {alias: 2031, 'valor_base': 0},
            })[('OEI.01', 0)]
            normalized.append(meta)

        self.assertEqual(normalized[0], normalized[1])
        self.assertEqual(normalized[0]['anio_base'], 2031)
        self.assertNotIn('año_base', normalized[0])

    def test_ficha_baseline_uses_form_values_atomically_and_falls_back(self):
        ficha = {'año_base': '2024', 'valor_absoluto': 'json-value'}

        overridden = resolver_linea_base_ficha(
            ficha,
            {'anio_base': 2031, 'valor_base': 0},
        )
        self.assertEqual(overridden['anio_base'], 2031)
        self.assertEqual(overridden['valor_absoluto'], 0)

        for incomplete_meta in (
            {'anio_base': 2031, 'valor_base': ''},
            {'anio_base': '', 'valor_base': 55},
            {},
        ):
            with self.subTest(meta=incomplete_meta):
                fallback = resolver_linea_base_ficha(ficha, incomplete_meta)
                self.assertEqual(fallback['anio_base'], '2024')
                self.assertEqual(fallback['valor_absoluto'], 'json-value')

    def test_fichas_keep_ordinal_and_all_annual_records(self):
        years = [2024, 2026, 2027, 2028, 2029, 2030]
        records = []
        for ordinal in range(2):
            for year in years:
                records.append({
                    'linea_base': 'Línea de Base' if year == 2024 else 'Logros esperados',
                    'anio_base': str(year),
                    'marker': f'{ordinal}-{year}',
                })

        indexed = indexar_fichas({'AEI.02.01': records})

        self.assertEqual(
            indexed[('AEI.02.01', 1, 2030)]['marker'],
            '1-2030',
        )
        selected = construir_fichas_seleccionadas(
            ['AEI.02.01'],
            {},
            {'AEI.02.01': {1}},
            indexed,
        )
        self.assertEqual(selected, [('AEI.02.01', 1)])

    def test_payload_without_fichas_does_not_create_ficha_entries(self):
        selected = construir_fichas_seleccionadas(
            ['OEI.01'],
            {'OEI.01': {0}},
            {},
            {},
        )

        self.assertEqual(selected, [])

    def test_period_years_use_real_template_columns(self):
        columns = mapear_columnas_anio(['Año (*):', '2024', '2026', '2027', '2028', '2029', '2030'])

        usable, missing = resolver_anios_columnas('2024-2030', columns)

        self.assertEqual(usable, [2024, 2026, 2027, 2028, 2029, 2030])
        self.assertEqual(missing, [2025])
        for period, expected in (
            ('2027-2032', [2027, 2028, 2029, 2030, 2031, 2032]),
            ('2024-2029', [2024, 2025, 2026, 2027, 2028, 2029]),
        ):
            with self.subTest(period=period):
                self.assertEqual(extraer_anios_periodo(period), expected)

    def test_baseline_is_selected_from_contract_not_period_start(self):
        records = []
        for year in (2024, 2026, 2027, 2028, 2029, 2030):
            records.append({
                'linea_base': 'Línea de Base' if year == 2024 else 'Logros esperados',
                'anio_base': str(year),
                'marker': f'{year}',
            })
        indexed = indexar_fichas({'OEI.01': records})

        for period in ('2027-2032', '2024-2029'):
            with self.subTest(period=period):
                baseline = seleccionar_ficha_base(
                    indexed,
                    'OEI.01',
                    0,
                    extraer_anios_periodo(period),
                )
                self.assertEqual(baseline['marker'], '2024')
                self.assertEqual(extraer_anios_periodo(period)[0], int(period[:4]))

    def test_period_years_expand_fixed_table_without_truncation(self):
        from docx import Document

        document = Document('PEI_Estandar_-_Informe.docx')
        table = document.tables[6]
        columns = ajustar_columnas_anio(table, [2027, 2028, 2029, 2030, 2031, 2032], 1, 5)

        self.assertEqual(list(columns), [2027, 2028, 2029, 2030, 2031, 2032])
        self.assertEqual(
            [cell.text for cell in table.rows[1].cells][5:],
            ['2027', '2028', '2029', '2030', '2031', '2032'],
        )
        self.assertEqual(len(table.rows[2].cells), 11)


class XmlCleanupTests(unittest.TestCase):
    def test_ordinance_placeholder_is_replaced_by_only_the_entered_value(self):
        from docx import Document

        document = Document()
        document.add_paragraph('Ordenanza N.º «Ordenanza_PDC»')
        reemplazar_placeholder_docx(document, '«Ordenanza_PDC»', '002-2024-MDL')

        self.assertEqual(document.paragraphs[0].text, 'Ordenanza N.º 002-2024-MDL')
        self.assertNotIn('«Ordenanza_PDC»', document.paragraphs[0].text)

    def test_split_placeholder_replacement_preserves_run_properties(self):
        from docx import Document
        from docx.shared import RGBColor

        document = Document()
        paragraph = document.add_paragraph()
        first = paragraph.add_run('«Ordenanza_')
        first.bold = True
        first.font.color.rgb = RGBColor(0x12, 0x34, 0x56)
        second = paragraph.add_run('PDC»')
        second.italic = True
        second.font.color.rgb = RGBColor(0x65, 0x43, 0x21)

        reemplazar_placeholder_docx(document, '«Ordenanza_PDC»', '002-2024-MDL')

        self.assertEqual(paragraph.text, '002-2024-MDL')
        self.assertTrue(first.bold)
        self.assertEqual(first.font.color.rgb, RGBColor(0x12, 0x34, 0x56))
        self.assertTrue(second.italic)
        self.assertEqual(second.font.color.rgb, RGBColor(0x65, 0x43, 0x21))

    def test_placeholder_paragraph_is_removed_without_touching_sectpr(self):
        namespace = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        body = ET.fromstring(
            f'<w:body xmlns:w="{namespace}">'
            '<w:p><w:r><w:t>Agregar fichas de indicadores</w:t></w:r></w:p>'
            '<w:p />'
            '<w:sectPr />'
            '</w:body>'
        )

        self.assertEqual(eliminar_parrafos_con_texto(body, 'Agregar fichas de indicadores'), 1)
        self.assertEqual(limpiar_parrafos_body_vacios_finales(body), 1)
        self.assertEqual(body[-1].tag.rsplit('}', 1)[-1], 'sectPr')

    def test_page_break_before_is_not_duplicated(self):
        from docx import Document

        paragraph = Document().add_paragraph()._p

        asegurar_salto_de_pagina_antes(paragraph)
        asegurar_salto_de_pagina_antes(paragraph)

        self.assertEqual(
            len(paragraph.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pageBreakBefore')),
            1,
        )

    def test_removes_only_trailing_empty_body_paragraphs(self):
        namespace = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        body = ET.fromstring(
            f'<w:body xmlns:w="{namespace}">'
            '<w:tbl />'
            '<w:p />'
            '<w:p><w:pPr /></w:p>'
            '<w:sectPr />'
            '</w:body>'
        )

        removed = limpiar_parrafos_body_vacios_finales(body)

        self.assertEqual(removed, 2)
        self.assertEqual([child.tag.rsplit("}", 1)[-1] for child in body], ['tbl', 'sectPr'])

    def test_preserves_page_break_paragraph(self):
        namespace = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        body = ET.fromstring(
            f'<w:body xmlns:w="{namespace}">'
            '<w:p><w:r><w:br w:type="page" /></w:r></w:p>'
            '<w:sectPr />'
            '</w:body>'
        )

        removed = limpiar_parrafos_body_vacios_finales(body)

        self.assertEqual(removed, 0)
        self.assertEqual(body[0].tag.rsplit("}", 1)[-1], 'p')

    def test_ficha_note_is_removed_as_a_row_and_preserved_as_paragraph(self):
        from docx import Document

        table = Document('Plantilla_de_ficha_técnica.docx').tables[0]._element
        note = extraer_nota_ficha(table)

        self.assertIsNotNone(note)
        self.assertEqual(len(table.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr')), 14)
        self.assertIn('Nota:', ''.join(node.text or '' for node in note.iter() if node.tag.rsplit('}', 1)[-1] == 't'))

    def test_t0_borders_do_not_contain_nil(self):
        from docx import Document

        document = Document('PEI_Estandar_-_Informe.docx')
        normalizar_bordes_tabla(document.tables[0])
        nil_borders = [
            element
            for element in document.tables[0]._element.iter()
            if element.tag.rsplit('}', 1)[-1] in {'top', 'left', 'bottom', 'right'}
            and element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') == 'nil'
        ]

        self.assertEqual(nil_borders, [])


class TemporaryFileCleanupTests(unittest.TestCase):
    def test_generation_wrapper_cleans_temporary_file_on_error(self):
        handler = object.__new__(PEIHandler)
        created = {}

        def fail(_data, temp_path):
            created['path'] = temp_path
            with open(temp_path, 'w', encoding='utf-8') as temp_file:
                temp_file.write('temporary')
            raise RuntimeError('expected')

        handler._generar_documento = fail
        with self.assertRaises(RuntimeError):
            handler.generar_documento({})

        self.assertFalse(os.path.exists(created['path']))


WORD_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
WORKTREE = Path(__file__).resolve().parent


def _local_name(element):
    return element.tag.rsplit('}', 1)[-1]


def _xml_text(element):
    return ''.join(
        child.text or ''
        for child in element.iter()
        if _local_name(child) == 't'
    )


def _table_rows(table):
    return table.findall(f'{{{WORD_NS}}}tr')


def _table_cells(row):
    return row.findall(f'{{{WORD_NS}}}tc')


def _empty_body_paragraph(element):
    for child in element.iter():
        name = _local_name(child)
        if name in {'p', 'pPr', 'r', 'rPr'}:
            continue
        if name == 't':
            if (child.text or '').strip():
                return False
            continue
        return False
    return True


@contextlib.contextmanager
def _mailmerge_adapter():
    from docx import Document

    class MailMergeFallback:
        def __init__(self, path):
            self.document = Document(path)

        def __enter__(self):
            return self

        def __exit__(self, *_args):
            return False

        def merge(self, **fields):
            for name, value in fields.items():
                reemplazar_placeholder_docx(
                    self.document,
                    f'«{name}»',
                    str(value),
                )

        def write(self, path):
            self.document.save(path)

    module = types.ModuleType('mailmerge')
    module.MailMerge = MailMergeFallback
    with patch.dict(sys.modules, {'mailmerge': module}):
        yield


def _fresh_payload(periodo):
    return {
        'codigo_ue': '150101',
        'nombre_municipio': 'Municipio Prueba',
        'periodo_pei': periodo,
        'ordenanza_pdc': '002-2024-MDL',
        'selecciones': {
            'oei': ['oei-OEI.01'],
            'aei': ['aei-AEI.01.01'],
            'indicadoresOEI': ['ind-oei-OEI.01-0'],
            'indicadoresAEI': ['ind-aei-AEI.01.01-0'],
        },
        'prioridades': {
            'oei': ['OEI.01'],
            'aei': {'OEI.01': ['AEI.01.01']},
        },
        'metas': {
            'ind-oei-OEI.01-0': {
                'año_base': 2024,
                'valor_base': 10,
                'meta_2027': 20,
                'meta_2032': 32,
            },
            'ind-aei-AEI.01.01-0': {
                'año_base': 2024,
                'valor_base': 2,
                'meta_2027': 3,
                'meta_2032': 8,
            },
        },
    }


def _generate_document_xml(data):
    with tempfile.TemporaryDirectory() as temp_dir:
        output_path = Path(temp_dir) / 'generated.docx'
        handler = object.__new__(PEIHandler)
        with _mailmerge_adapter():
            handler.generar_documento(
                data,
                template_dir=WORKTREE,
                output_path=output_path,
            )
        with zipfile.ZipFile(output_path) as archive:
            return ET.fromstring(archive.read('word/document.xml'))


def _ordered_generation_payload():
    indicator_ids = [
        'ind-oei-OEI.01-0',
        'ind-oei-OEI.02-0',
        'ind-aei-AEI.01.01-0',
        'ind-aei-AEI.01.02-0',
        'ind-aei-AEI.02.01-0',
    ]
    return {
        'codigo_ue': '150101',
        'nombre_municipio': 'Municipio Ordenado',
        'periodo_pei': '2026-2030',
        'ordenanza_pdc': '002-2024-MDL',
        'selecciones': {
            'oei': ['oei-OEI.01', 'oei-OEI.02'],
            'aei': ['aei-AEI.01.01', 'aei-AEI.01.02', 'aei-AEI.02.01'],
            'indicadoresOEI': indicator_ids[:2],
            'indicadoresAEI': indicator_ids[2:],
        },
        'prioridades': {
            'oei': ['OEI.02', 'OEI.01'],
            'aei': {
                'OEI.02': ['AEI.02.01'],
                'OEI.01': ['AEI.01.02', 'AEI.01.01'],
            },
        },
        'metas': {
            indicator_id: {
                'año_base': 2024,
                'valor_base': index + 1,
                'meta_2026': index + 10,
                'meta_2027': index + 20,
                'meta_2030': index + 30,
            }
            for index, indicator_id in enumerate(indicator_ids)
        },
    }


def _catalog_display_payload():
    indicator_ids = [
        'ind-oei-OEI.01-0',
        'ind-oei-OEI.04-0',
        'ind-oei-OEI.10-0',
        'ind-aei-AEI.01.05-0',
        'ind-aei-AEI.04.02-0',
        'ind-aei-AEI.10.01-0',
    ]
    return {
        'codigo_ue': '150101',
        'nombre_municipio': 'Municipio Catalogo',
        'periodo_pei': '2026-2030',
        'ordenanza_pdc': '002-2024-MDL',
        'selecciones': {
            'oei': ['oei-OEI.01', 'oei-OEI.04', 'oei-OEI.10'],
            'aei': ['aei-AEI.01.05', 'aei-AEI.04.02', 'aei-AEI.10.01'],
            'indicadoresOEI': indicator_ids[:3],
            'indicadoresAEI': indicator_ids[3:],
        },
        'prioridades': {
            'oei': ['OEI.01', 'OEI.04', 'OEI.10'],
            'aei': {
                'OEI.01': ['AEI.01.05'],
                'OEI.04': ['AEI.04.02'],
                'OEI.10': ['AEI.10.01'],
            },
        },
        'metas': {
            indicator_id: {
                'año_base': 2024,
                'valor_base': index + 1,
                'meta_2026': index + 10,
                'meta_2027': index + 20,
                'meta_2030': index + 30,
            }
            for index, indicator_id in enumerate(indicator_ids)
        },
    }


def _partial_indicator_payload():
    return {
        'codigo_ue': '150101',
        'nombre_municipio': 'Municipio Parcial',
        'periodo_pei': '2026-2030',
        'ordenanza_pdc': '002-2024-MDL',
        'selecciones': {
            'oei': ['oei-OEI.02'],
            'aei': ['aei-AEI.02.01'],
            'indicadoresOEI': [],
            'indicadoresAEI': ['ind-aei-AEI.02.01-1'],
        },
        'prioridades': {
            'oei': ['OEI.02'],
            'aei': {'OEI.02': ['AEI.02.01']},
        },
        'metas': {
            'ind-aei-AEI.02.01-0': {
                'año_base': 2024,
                'valor_base': 90,
                'meta_2026': 900,
                'meta_2027': 901,
            },
            'ind-aei-AEI.02.01-1': {
                'año_base': 2024,
                'valor_base': 101,
                'meta_2026': 201,
                'meta_2027': 202,
                'meta_2030': 205,
            },
        },
    }


def _has_page_break_before(paragraph):
    p_pr = paragraph.find(f'{{{WORD_NS}}}pPr')
    return p_pr is not None and p_pr.find(f'{{{WORD_NS}}}pageBreakBefore') is not None


class FreshDocxGenerationTests(unittest.TestCase):
    def test_payload_baseline_overrides_json_in_matrix_and_ficha(self):
        data = _fresh_payload('2027-2032')
        data['metas']['ind-oei-OEI.01-0'].update({
            'año_base': 2031,
            'valor_base': 901.5,
        })
        data['metas']['ind-aei-AEI.01.01-0'].update({
            'anio_base': 2032,
            'valor_base': 0,
        })

        root = _generate_document_xml(data)
        tables = root.findall(f'.//{{{WORD_NS}}}tbl')
        matrix = next(
            table for table in tables
            if 'Código' in _xml_text(table)
            and 'Logros Esperados' in _xml_text(table)
        )
        expected = {
            'OEI.01': ('2031', '901.5'),
            'AEI.01.01': ('2032', '0'),
        }
        for code, baseline in expected.items():
            row = next(
                row for row in _table_rows(matrix)
                if _xml_text(_table_cells(row)[0]).strip() == code
            )
            cells = _table_cells(row)
            self.assertEqual((_xml_text(cells[3]), _xml_text(cells[4])), baseline)

        ficha_tables = [
            table for table in tables
            if 'FICHA TÉCNICA DEL INDICADOR' in _xml_text(table)
        ]
        for table in ficha_tables:
            code = re.search(
                r'Objetivo/Acción:(OEI\.\d{2}|AEI\.\d{2}\.\d{2})',
                _xml_text(table),
            ).group(1)
            baseline_row = _table_rows(table)[11]
            absolute_row = _table_rows(table)[13]
            self.assertEqual(
                (_xml_text(_table_cells(baseline_row)[1]),
                 _xml_text(_table_cells(absolute_row)[1])),
                expected[code],
            )

    def test_same_aei_indicator_ordinals_keep_their_baselines(self):
        data = _partial_indicator_payload()
        data['selecciones']['indicadoresAEI'] = [
            'ind-aei-AEI.02.01-0',
            'ind-aei-AEI.02.01-1',
        ]
        data['metas']['ind-aei-AEI.02.01-0'].update({
            'año_base': 2031,
            'valor_base': 0,
        })
        data['metas']['ind-aei-AEI.02.01-1'].update({
            'anio_base': 2032,
            'valor_base': 222,
        })

        root = _generate_document_xml(data)
        tables = root.findall(f'.//{{{WORD_NS}}}tbl')
        matrix = next(
            table for table in tables
            if 'OEI/AEI' in _xml_text(table)
            and 'Logros Esperados' in _xml_text(table)
        )
        rows = [
            row for row in _table_rows(matrix)
            if _table_cells(row)
            and _xml_text(_table_cells(row)[0]).strip() == 'AEI.01.01'
        ]
        self.assertEqual(
            [
                (_xml_text(_table_cells(row)[3]), _xml_text(_table_cells(row)[4]))
                for row in rows
            ],
            [('2031', '0'), ('2032', '222')],
        )

        ficha_tables = [
            table for table in tables
            if 'FICHA TÉCNICA DEL INDICADOR' in _xml_text(table)
        ]
        self.assertEqual(len(ficha_tables), 2)
        self.assertEqual(
            [
                (
                    _xml_text(_table_cells(_table_rows(table)[11])[1]),
                    _xml_text(_table_cells(_table_rows(table)[13])[1]),
                )
                for table in ficha_tables
            ],
            [('2031', '0'), ('2032', '222')],
        )

    def test_fresh_docx_preserves_period_selection_and_xml_contract(self):
        for period in ('2027-2032', '2024-2029'):
            with self.subTest(period=period), tempfile.TemporaryDirectory() as temp_dir:
                output_path = Path(temp_dir) / f'fresh-{period}.docx'
                handler = object.__new__(PEIHandler)

                with _mailmerge_adapter():
                    result = handler.generar_documento(
                        _fresh_payload(period),
                        template_dir=WORKTREE,
                        output_path=output_path,
                    )

                self.assertEqual(result, str(output_path))
                self.assertTrue(output_path.is_file())
                with zipfile.ZipFile(output_path) as archive:
                    root = ET.fromstring(archive.read('word/document.xml'))

                full_text = _xml_text(root)
                self.assertIn('002-2024-MDL', full_text)
                self.assertNotIn('«Ordenanza_PDC»', full_text)
                self.assertNotIn('Agregar fichas de indicadores', full_text)
                for unselected in ('OEI.02', 'AEI.01.02', 'AEI.02.01'):
                    self.assertNotIn(unselected, full_text)

                tables = root.findall(f'.//{{{WORD_NS}}}tbl')
                matrix = next(
                    table for table in tables
                    if 'Código' in _xml_text(table)
                    and 'Logros Esperados' in _xml_text(table)
                )
                matrix_header = next(
                    row for row in _table_rows(matrix)
                    if 'Código' in _xml_text(row)
                )
                matrix_years = [
                    _xml_text(cell)
                    for cell in _table_cells(matrix_header)
                    if _xml_text(cell).isdigit()
                ]
                self.assertEqual(
                    matrix_years,
                    [str(year) for year in extraer_anios_periodo(period)],
                )
                expected_meta = {
                    'OEI.01': {'2027': '20', '2032': '32'},
                    'AEI.01.01': {'2027': '3', '2032': '8'},
                }
                for code in expected_meta:
                    data_row = next(
                        row for row in _table_rows(matrix)
                        if _xml_text(_table_cells(row)[0]) == code
                    )
                    values = {
                        year: _xml_text(cell)
                        for year, cell in zip(
                            matrix_years,
                            _table_cells(data_row)[5:],
                        )
                    }
                    for year in matrix_years:
                        self.assertEqual(values[year], expected_meta[code].get(year, ''))

                ficha_tables = [
                    table for table in tables
                    if 'FICHA TÉCNICA DEL INDICADOR' in _xml_text(table)
                ]
                self.assertEqual(len(ficha_tables), 2)
                expected_years = [str(year) for year in extraer_anios_periodo(period)]
                for ficha_table in ficha_tables:
                    year_row = _table_rows(ficha_table)[11]
                    cells = [_xml_text(cell) for cell in _table_cells(year_row)]
                    self.assertEqual(cells[1], '2024')
                    self.assertEqual(cells[2:], expected_years)
                    absolute_row = _table_rows(ficha_table)[13]
                    self.assertNotIn('Nota:', _xml_text(absolute_row))
                    achievement_values = [
                        _xml_text(cell)
                        for cell in _table_cells(_table_rows(ficha_table)[12])[2:]
                    ]
                    for year, value in zip(expected_years, achievement_values):
                        if year in {'2025', '2031', '2032'}:
                            self.assertEqual(value, '')

                body = root.find(f'{{{WORD_NS}}}body')
                external_notes = [
                    child for child in body
                    if _local_name(child) == 'p' and 'Nota:' in _xml_text(child)
                ]
                self.assertEqual(len(external_notes), len(ficha_tables))
                self.assertEqual(_local_name(body[-1]), 'sectPr')
                self.assertEqual(
                    sum(1 for child in body if _local_name(child) == 'sectPr'),
                    1,
                )
                self.assertFalse(
                    _local_name(body[-2]) == 'p' and _empty_body_paragraph(body[-2])
                )

                nil_borders = [
                    element for element in tables[0].iter()
                    if _local_name(element) in {'top', 'left', 'bottom', 'right'}
                    and element.get(f'{{{WORD_NS}}}val') == 'nil'
                ]
                self.assertEqual(nil_borders, [])

                self.assertEqual(
                    len(list(Path(temp_dir).glob('*.docx'))),
                    1,
                )

    def test_anexo_b1_typography_is_scoped_to_target_sentence(self):
        from docx import Document
        from docx.shared import Pt

        with tempfile.TemporaryDirectory() as template_dir, tempfile.TemporaryDirectory() as output_dir:
            template_dir = Path(template_dir)
            shutil.copy(WORKTREE / 'PEI_Estandar_-_Informe.docx', template_dir)
            shutil.copy(WORKTREE / 'Plantilla_de_ficha_técnica.docx', template_dir)

            template_path = template_dir / 'PEI_Estandar_-_Informe.docx'
            template = Document(template_path)
            control = next(
                paragraph for paragraph in template.paragraphs
                if paragraph.text.startswith('ANEXO B - 2')
            )
            control.runs[0].font.name = 'Calibri'
            control.runs[0].font.size = Pt(9)
            template.save(template_path)

            output_path = Path(output_dir) / 'generated.docx'
            handler = object.__new__(PEIHandler)
            with _mailmerge_adapter():
                handler.generar_documento(
                    _fresh_payload('2027-2032'),
                    template_dir=template_dir,
                    output_path=output_path,
                )

            with zipfile.ZipFile(output_path) as archive:
                root = ET.fromstring(archive.read('word/document.xml'))

        body = root.find(f'{{{WORD_NS}}}body')
        paragraphs = body.findall(f'{{{WORD_NS}}}p')
        target = next(
            paragraph for paragraph in paragraphs
            if _xml_text(paragraph).startswith('El PEI está articulado al')
        )
        target_runs = target.findall(f'.//{{{WORD_NS}}}r')
        self.assertTrue(target_runs)
        for run in target_runs:
            r_pr = run.find(f'{{{WORD_NS}}}rPr')
            r_fonts = r_pr.find(f'{{{WORD_NS}}}rFonts')
            self.assertEqual(
                {name: r_fonts.get(f'{{{WORD_NS}}}{name}') for name in ('ascii', 'eastAsia', 'hAnsi', 'cs')},
                {name: 'Arial Narrow' for name in ('ascii', 'eastAsia', 'hAnsi', 'cs')},
            )
            self.assertEqual(r_pr.find(f'{{{WORD_NS}}}sz').get(f'{{{WORD_NS}}}val'), '24')
            self.assertEqual(r_pr.find(f'{{{WORD_NS}}}szCs').get(f'{{{WORD_NS}}}val'), '24')

        self.assertIn(
            'FF0000',
            {
                run.find(f'{{{WORD_NS}}}rPr/{{{WORD_NS}}}color').get(f'{{{WORD_NS}}}val')
                for run in target_runs
                if run.find(f'{{{WORD_NS}}}rPr/{{{WORD_NS}}}color') is not None
            },
        )

        control = next(
            paragraph for paragraph in paragraphs
            if _xml_text(paragraph).startswith('ANEXO B - 2')
        )
        control_r_pr = control.find(f'{{{WORD_NS}}}r/{{{WORD_NS}}}rPr')
        self.assertEqual(
            control_r_pr.find(f'{{{WORD_NS}}}rFonts').get(f'{{{WORD_NS}}}ascii'),
            'Calibri',
        )
        self.assertEqual(
            control_r_pr.find(f'{{{WORD_NS}}}sz').get(f'{{{WORD_NS}}}val'),
            '18',
        )
        self.assertEqual(
            control.find(f'{{{WORD_NS}}}r/{{{WORD_NS}}}rPr/{{{WORD_NS}}}b') is not None,
            True,
        )

    def test_generation_wrapper_cleans_temporary_file_on_success(self):
        handler = object.__new__(PEIHandler)
        created = {}

        def succeed(_data, temp_path):
            created['path'] = temp_path
            return 'output.docx'

        handler._generar_documento = succeed

        self.assertEqual(handler.generar_documento({}), 'output.docx')
        self.assertFalse(os.path.exists(created['path']))

    def test_catalog_codes_are_remapped_only_in_visible_docx_xml(self):
        data = _catalog_display_payload()
        root = _generate_document_xml(data)
        full_text = _xml_text(root)
        visible_codes = set(re.findall(
            r'(?<![A-Za-z0-9_.-])(OEI\.\d{2}|AEI\.\d{2}\.\d{2})(?![A-Za-z0-9_.-])',
            full_text,
        ))
        display_codes = {
            'OEI.01', 'OEI.02', 'OEI.03',
            'AEI.01.01', 'AEI.02.01', 'AEI.03.01',
        }

        self.assertTrue(visible_codes)
        self.assertTrue(visible_codes <= display_codes)
        for original in ('OEI.04', 'OEI.10', 'AEI.01.05', 'AEI.04.02', 'AEI.10.01'):
            self.assertNotIn(original, full_text)
        for display in display_codes:
            self.assertIn(display, full_text)

        tables = root.findall(f'.//{{{WORD_NS}}}tbl')
        for table in tables:
            table_codes = set(re.findall(
                r'(?<![A-Za-z0-9_.-])(OEI\.\d{2}|AEI\.\d{2}\.\d{2})(?![A-Za-z0-9_.-])',
                _xml_text(table),
            ))
            self.assertTrue(table_codes <= display_codes)

        self.assertEqual(data['selecciones']['oei'], [
            'oei-OEI.01', 'oei-OEI.04', 'oei-OEI.10'
        ])
        self.assertEqual(data['prioridades']['oei'], ['OEI.01', 'OEI.04', 'OEI.10'])
        self.assertEqual(data['prioridades']['aei']['OEI.01'], ['AEI.01.05'])
        self.assertEqual(data['selecciones']['aei'][0], 'aei-AEI.01.05')
        self.assertIn(('AEI.01.05', 0), indexar_metas(data['metas']))
        self.assertIn(('AEI.04.02', 0), indexar_metas(data['metas']))

    def test_two_objectives_keep_priority_order_across_route_matrix_and_fichas(self):
        root = _generate_document_xml(_ordered_generation_payload())
        full_text = _xml_text(root)
        tables = root.findall(f'.//{{{WORD_NS}}}tbl')

        route = next(
            table for table in tables
            if 'Cod_OEI' in _xml_text(table) and 'Cod AEI' in _xml_text(table)
        )
        route_rows = []
        for row in _table_rows(route):
            cells = _table_cells(row)
            if len(cells) > 5 and re.fullmatch(r'OEI\.\d{2}', _xml_text(cells[1]).strip()):
                route_rows.append([
                    _xml_text(cells[index]).strip()
                    for index in (0, 1, 4, 5)
                ])
        self.assertEqual(route_rows, [
            ['1', 'OEI.01', '1', 'AEI.01.01'],
            ['2', 'OEI.02', '1', 'AEI.02.01'],
            ['2', 'OEI.02', '2', 'AEI.02.02'],
        ])

        matrix = next(
            table for table in tables
            if 'OEI/AEI' in _xml_text(table) and 'Logros Esperados' in _xml_text(table)
        )
        matrix_codes = [
            _xml_text(_table_cells(row)[0]).strip()
            for row in _table_rows(matrix)
            if _table_cells(row)
            and re.fullmatch(r'(?:OEI\.\d{2}|AEI\.\d{2}\.\d{2})', _xml_text(_table_cells(row)[0]).strip())
        ]
        self.assertEqual(matrix_codes, [
            'OEI.02', 'AEI.02.02', 'AEI.02.01', 'OEI.01', 'AEI.01.01'
        ])

        ficha_tables = [
            table for table in tables
            if 'FICHA TÉCNICA DEL INDICADOR' in _xml_text(table)
        ]
        ficha_codes = [
            re.search(
                r'Objetivo/Acción:(OEI\.\d{2}|AEI\.\d{2}\.\d{2})',
                _xml_text(table),
            ).group(1)
            for table in ficha_tables
        ]
        self.assertEqual(ficha_codes, [
            'OEI.01', 'AEI.01.01', 'OEI.02', 'AEI.02.01', 'AEI.02.02'
        ])

        for code in ('OEI.03', 'OEI.04', 'OEI.05', 'OEI.06', 'OEI.07', 'OEI.08', 'OEI.09', 'OEI.10', 'OEI.11'):
            self.assertNotIn(
                f'Acciones Estratégicas Institucionales del {code}',
                full_text,
            )

        body = root.find(f'{{{WORD_NS}}}body')
        external_notes = [
            child for child in body
            if _local_name(child) == 'p' and 'Nota:' in _xml_text(child)
        ]
        self.assertEqual(len(external_notes), len(ficha_tables))
        for table in ficha_tables:
            table_index = list(body).index(table)
            note = body[table_index + 1]
            self.assertEqual(_local_name(note), 'p')
            self.assertIn('Nota:', _xml_text(note))
            self.assertNotIn('Nota:', _xml_text(_table_rows(table)[13]))
            self.assertFalse(any(
                'Nota:' in _xml_text(cell)
                for row in _table_rows(table)
                for cell in _table_cells(row)
            ))

        self.assertNotIn('«Ordenanza_PDC»', full_text)
        self.assertNotIn('Agregar fichas de indicadores', full_text)
        self.assertEqual(_local_name(body[-1]), 'sectPr')
        self.assertEqual(
            sum(1 for child in body if _local_name(child) == 'sectPr'),
            1,
        )
        self.assertFalse(_empty_body_paragraph(body[-2]))

        anexo_title = next(
            child for child in body
            if _local_name(child) == 'p' and 'ANEXO A - 6' in _xml_text(child)
        )
        self.assertTrue(_has_page_break_before(anexo_title))
        for index, table in enumerate(ficha_tables):
            first_paragraph = table.find(f'{{{WORD_NS}}}tr/{{{WORD_NS}}}tc/{{{WORD_NS}}}p')
            self.assertEqual(_has_page_break_before(first_paragraph), index > 0)
        self.assertEqual(
            len(root.findall(f'.//{{{WORD_NS}}}pageBreakBefore')),
            len(ficha_tables),
        )

    def test_partial_indicator_ordinal_keeps_only_its_meta_and_ficha(self):
        root = _generate_document_xml(_partial_indicator_payload())
        full_text = _xml_text(root)
        selected_name = 'Porcentaje de inversiones en infraestructura culminadas para la mejora/construcción de establecimientos de salud bajo administración municipal'
        unselected_name = 'Porcentaje de establecimientos médicos bajo responsabilidad municipal en condiciones operativas y de infraestructura adecuadas'
        self.assertIn(selected_name, full_text)
        self.assertNotIn(unselected_name, full_text)

        tables = root.findall(f'.//{{{WORD_NS}}}tbl')
        matrix = next(
            table for table in tables
            if 'OEI/AEI' in _xml_text(table) and 'Logros Esperados' in _xml_text(table)
        )
        rows = [
            row for row in _table_rows(matrix)
            if _table_cells(row)
            and _xml_text(_table_cells(row)[0]).strip() == 'AEI.01.01'
        ]
        self.assertEqual(len(rows), 1)
        cells = _table_cells(rows[0])
        self.assertEqual(_xml_text(cells[3]).strip(), '2024')
        self.assertEqual(_xml_text(cells[4]).strip(), '101')
        self.assertEqual(
            [_xml_text(cell).strip() for cell in cells[5:]],
            ['201', '202', '', '', '205'],
        )
        self.assertNotIn('900', _xml_text(rows[0]))

        ficha_tables = [
            table for table in tables
            if 'FICHA TÉCNICA DEL INDICADOR' in _xml_text(table)
        ]
        self.assertEqual(len(ficha_tables), 1)
        ficha = ficha_tables[0]
        self.assertIn(selected_name, _xml_text(ficha))
        self.assertNotIn(unselected_name, _xml_text(ficha))
        year_cells = [_xml_text(cell).strip() for cell in _table_cells(_table_rows(ficha)[11])]
        self.assertEqual(year_cells[1:], ['2024', '2026', '2027', '2028', '2029', '2030'])
        self.assertNotIn('Nota:', _xml_text(_table_rows(ficha)[13]))


class SecurityRegressionTests(unittest.TestCase):
    def test_sys_is_imported_for_handler_error_logging(self):
        self.assertIs(generador_pei_fase4.sys, sys)

    def test_output_name_is_sanitized(self):
        self.assertEqual(
            sanitizar_nombre_municipio('Municipalidad Provincial'),
            'Municipalidad_Provincial',
        )
        with self.assertRaises(ValueError):
            sanitizar_nombre_municipio('../salida')

    def test_static_allowlist_does_not_expose_source_files(self):
        handler = object.__new__(PEIHandler)
        self.assertIsNone(handler._allowed_static_file('/generador_pei_fase4.py'))


if __name__ == '__main__':
    unittest.main()
