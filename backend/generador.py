# Tus importaciones actuales al inicio de generador.py
from http.server import HTTPServer, SimpleHTTPRequestHandler
from functools import partial
import json, os, sys, re, io, urllib.parse, subprocess, importlib.util
from pathlib import Path
from urllib.parse import urlparse
import pandas as pd
import warnings
warnings.filterwarnings('ignore', category=UserWarning, module='openpyxl')

# --- AÑADE ESTAS LÍNEAS AQUÍ ARRIBA ---
from mailmerge import MailMerge
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK
from copy import deepcopy
# -------------------------------------

PROJECT_ROOT = Path(__file__).resolve().parent.parent
FRONTEND_DIR = PROJECT_ROOT / 'frontend'
PLANTILLAS_DIR = PROJECT_ROOT / 'plantillas'
EXCEL_PATH = PROJECT_ROOT / 'IT_PEI.xlsx'

# Almacén global en memoria para las Unidades Ejecutoras
DATOS_UES = {}

# --- NUEVO: Cargar variables desde .env.local (Solo para uso local) ---
env_path = PROJECT_ROOT / '.env.local'
if env_path.exists():
    with open(env_path, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            # Ignorar líneas vacías o comentarios
            if line and not line.startswith('#') and '=' in line:
                key, value = line.split('=', 1)
                os.environ[key.strip()] = value.strip()

def cargar_catalogo_ues():
    """Carga el catálogo de Unidades Ejecutoras desde AMBAS hojas del archivo Excel IT_PEI.xlsx."""
    global DATOS_UES
    
    excel_file = resolver_ruta(EXCEL_PATH, PROJECT_ROOT / 'IT_PEI.xlsx')
    
    if not excel_file.exists():
        print(f"⚠️ AVISO: No se encontró el archivo de datos Excel en: {excel_file}")
        return

    try:
        # 1. Leer hoja 'IT PEI' (contiene municipalidades como 300682)
        df_it = pd.read_excel(excel_file, sheet_name='IT PEI')
        for _, row in df_it.iterrows():
            id_ue = str(row.get('Id_UE', '')).strip()
            if id_ue.endswith('.0'):
                id_ue = id_ue[:-2]
                
            if id_ue and id_ue != 'nan':
                DATOS_UES[id_ue] = {
                    "codigo_ue": id_ue,
                    "nombre_municipio": str(row.get('Nombre_unidad_ejecutora', '')).strip(),
                    "nombre_provincia": str(row.get('nombre_provincia', '')).strip(),
                    "nombre_region": str(row.get('nombre_departamento', '')).strip(),
                    "periodo_pei": str(row.get('Periodo PEI', '')).strip() if str(row.get('Periodo PEI', '')).strip() != 'nan' else ''
                }

        # 2. Leer hoja 'Data_UEs' (complementa o actualiza datos)
        df_ues = pd.read_excel(excel_file, sheet_name='Data_UEs')
        for _, row in df_ues.iterrows():
            id_ue = str(row.get('Id_UE', '')).strip()
            if id_ue.endswith('.0'):
                id_ue = id_ue[:-2]
                
            if id_ue and id_ue != 'nan':
                DATOS_UES[id_ue] = {
                    "codigo_ue": id_ue,
                    "nombre_municipio": str(row.get('Nombre_unidad_ejecutora', '')).strip(),
                    "nombre_provincia": str(row.get('nombre_provincia', '')).strip(),
                    "nombre_region": str(row.get('nombre_departamento', '')).strip(),
                    "periodo_pei": str(row.get('Ult.PEI', '')).strip() if str(row.get('Ult.PEI', '')).strip() != 'nan' else ''
                }

        print(f"✅ Catálogo de UEs cargado con éxito: {len(DATOS_UES)} entidades (incluye Municipalidades).")
    except Exception as e:
        print(f"⚠️ Error al cargar el archivo Excel IT_PEI.xlsx: {e}")


def resolver_ruta(*parts):
    """Devuelve la primera ruta existente entre varias opciones posibles."""
    candidates = []
    for part in parts:
        candidates.append(Path(part))
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return candidates[0]



class PEIHandler(SimpleHTTPRequestHandler):
    def __init__(self, *args, directory=None, **kwargs):
        super().__init__(*args, directory=str(directory or PROJECT_ROOT), **kwargs)

    def do_OPTIONS(self):
        # Permitir CORS (para cuando esté en producción)
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'GET, POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.end_headers()

    def do_GET(self):
        # API Endpoint para Autocompletar datos por Código UE
        path = urlparse(self.path).path
        if path.startswith('/api/ue'):
            id_buscado = path.rstrip('/').split('/')[-1].strip() if path.rstrip('/').count('/') >= 3 else ''
            self.send_response(200)
            self.send_header('Content-type', 'application/json; charset=utf-8')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.send_header('Access-Control-Allow-Methods', 'GET, OPTIONS')
            self.send_header('Access-Control-Allow-Headers', 'Content-Type')
            self.end_headers()

            if self.command == 'OPTIONS':
                return

            if id_buscado in DATOS_UES:
                respuesta = {"success": True, "data": DATOS_UES[id_buscado]}
            else:
                respuesta = {"success": False, "message": "Código UE no encontrado en el catálogo"}

            self.wfile.write(json.dumps(respuesta, ensure_ascii=False).encode('utf-8'))
            return

        return super().do_GET()

    def do_POST(self):
        if self.path == '/generar':
            try:
                data = json.loads(self.rfile.read(int(self.headers['Content-Length'])).decode('utf-8'))
                
                # Retorna los bytes y el nombre en lugar de una ruta de archivo
                file_bytes, file_name = self.generar_documento(data)
                
                self.send_response(200)
                self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                self.send_header('Content-Disposition', f'attachment; filename="{urllib.parse.quote(file_name)}"')
                self.send_header('Access-Control-Allow-Origin', '*')
                self.send_header('Access-Control-Expose-Headers', 'Content-Disposition')
                self.send_header('Content-Length', str(len(file_bytes)))
                self.end_headers()
                
                # Escribe los bytes directamente a la respuesta
                self.wfile.write(file_bytes)
            except Exception as e:
                print(f"ERROR: {e}")
                import traceback; traceback.print_exc()
                self.send_response(500)
                self.send_header('Content-type', 'application/json; charset=utf-8')
                self.send_header('Access-Control-Allow-Origin', '*')
                self.end_headers()
                self.wfile.write(json.dumps({'success': False, 'error': str(e)}, ensure_ascii=False).encode('utf-8'))
        else:
            self.send_error(404)

    def log_message(self, format, *args):
        pass

def generar_documento(self, data):
        print("\n" + "="*70)
        print("GENERANDO DOCUMENTO PEI EN MEMORIA")
        print("="*70)

        campos = ['codigo_ue', 'nombre_municipio', 'periodo_pei', 'nombre_provincia',
                  'nombre_region', 'nombre_alcalde', 'resolucion_alcaldia',
                  'plan_desarrollo_concertado', 'ordenanza_pdc', 'mision_institucional',
                  'politica_general_gobierno', 'decreto_politica_general_gobierno']
        
        d = {k: str(data.get(k, '')).strip() for k in campos}
        print(f"Municipio: {d['nombre_municipio']}")

        sel = data.get('selecciones', {})
        oei_list = [x.replace('oei-', '') for x in sel.get('oei', [])]
        aei_list = [x.replace('aei-', '') for x in sel.get('aei', [])]
        
        prioridades = data.get('prioridades', {})
        oei_priorizados = prioridades.get('oei', [])
        
        if oei_priorizados:
            oei_ordenados = [x.replace('oei-', '') for x in oei_priorizados if x.replace('oei-', '') in oei_list]
        else:
            oei_ordenados = oei_list

        codigos_ordenados = []
        for oei_code in oei_ordenados:
            codigos_ordenados.append(oei_code)
            oei_num = oei_code.split('.')[1]
            aei_de_este_oei = [
                aei for aei in aei_list 
                if aei.startswith(f'AEI.{oei_num}.')
            ]
            aei_de_este_oei.sort()
            codigos_ordenados.extend(aei_de_este_oei)

        metas_raw = data.get('metas', {})
        metas_por_codigo = {}
        for key, meta in metas_raw.items():
            m = re.search(r'ind-oei-(OEI\.\d+)-', key) or re.search(r'ind-aei-(AEI\.\d+\.\d+)-', key)
            if m:
                metas_por_codigo.setdefault(m.group(1), []).append(meta)

        anios_match = re.findall(r'\d{4}', d['periodo_pei'])
        anios = list(range(int(anios_match[0]), int(anios_match[1])+1)) if len(anios_match)>=2 else [2026,2027,2028,2029,2030]

        # PASO 1
        print("\nPASO 1: MailMerge en memoria...")
        merge_data = {
            'Codigo_UE': d['codigo_ue'], 'Nombre_Municipio': d['nombre_municipio'],
            'Periodo_PEI': d['periodo_pei'], 'Nombre_Provincia': d['nombre_provincia'],
            'Nombre_Region': d['nombre_region'], 'Nombre_Alcalde': d['nombre_alcalde'],
            'Resolucion_Alcaldia': d['resolucion_alcaldia'],
            'Plan_Desarrollo_Concertado': d['plan_desarrollo_concertado'],
            'Ordenanza_PDC': d['ordenanza_pdc'], 'Mision_Institucional': d['mision_institucional'],
            'Politica_General_Gobierno': d['politica_general_gobierno'],
            'Decreto_Politica_General_Gobierno': d['decreto_politica_general_gobierno'],
        }
        plantilla_informe = resolver_ruta(PLANTILLAS_DIR / 'PEI_Estandar_-_Informe.docx', PROJECT_ROOT / 'PEI_Estandar_-_Informe.docx')
        
        # Guardar el MailMerge en un buffer de memoria
        merge_buffer = io.BytesIO()
        with MailMerge(str(plantilla_informe)) as doc_merge:
            doc_merge.merge(**merge_data)
            doc_merge.write(merge_buffer)
        
        merge_buffer.seek(0)
        print("  OK")

        # PASO 2
        print("\nPASO 2: Filtrar...")
        doc = Document(merge_buffer)

        ind_oei_sel = sel.get('indicadoresOEI', [])
        ind_aei_sel = sel.get('indicadoresAEI', [])

        indices_oei, indices_aei = {}, {}
        for ind_id in ind_oei_sel:
            m = re.match(r'ind-oei-(OEI\.\d+)-(\d+)', ind_id)
            if m: indices_oei.setdefault(m.group(1), set()).add(int(m.group(2)))

        for ind_id in ind_aei_sel:
            m = re.match(r'ind-aei-(AEI\.\d+\.\d+)-(\d+)', ind_id)
            if m: indices_aei.setdefault(m.group(1), set()).add(int(m.group(2)))

        def filtrar_tabla(table, patron, lista, indices_map, col_codigo):
            fdel, contadores = [], {}
            for i, row in enumerate(table.rows):
                if len(row.cells) <= col_codigo: continue
                m = re.search(patron, row.cells[col_codigo].text.strip())
                if m:
                    codigo = m.group()
                    if codigo not in lista:
                        fdel.append(i)
                    elif codigo in indices_map:
                        idx = contadores.get(codigo, 0)
                        contadores[codigo] = idx + 1
                        if idx not in indices_map[codigo]: fdel.append(i)
            for i in reversed(fdel): table._element.remove(table.rows[i]._element)
            return len(fdel)

        def detectar_col(table, patron):
            for row in table.rows[1:4]:
                for c_idx, cell in enumerate(row.cells):
                    if re.search(patron, cell.text.strip()): return c_idx
            return -1

        def filtrar_tabla_combinada(table, col_oei, col_aei):
            fdel = []
            for i, row in enumerate(table.rows):
                ncols = len(row.cells)
                if ncols <= max(col_oei, col_aei): continue
                m_oei = re.search(r'OEI\.\d{2}', row.cells[col_oei].text.strip())
                m_aei = re.search(r'AEI\.\d{2}\.\d{2}', row.cells[col_aei].text.strip())
                if m_oei or m_aei:
                    oei_ok = (not m_oei) or (m_oei.group() in oei_list)
                    aei_ok = (not m_aei) or (m_aei.group() in aei_list)
                    if not oei_ok or not aei_ok: fdel.append(i)
            for i in reversed(fdel): table._element.remove(table.rows[i]._element)
            return len(fdel)

        def es_tabla_matriz(table):
            for row in table.rows[2:5]:
                if row.cells and re.search(r'OEI\.\d{2}', row.cells[0].text): return True
            return False

        def filtrar_tabla_matriz(table):
            fdel, contadores_oei, contadores_aei, oei_activo = [], {}, {}, False
            for i, row in enumerate(table.rows):
                if not row.cells: continue
                col0 = row.cells[0].text.strip()
                m_aei = re.search(r'AEI\.\d{2}\.\d{2}', col0)
                m_oei = re.search(r'OEI\.\d{2}', col0)

                if m_aei:
                    codigo_aei = m_aei.group()
                    if not oei_activo or codigo_aei not in aei_list: fdel.append(i)
                    elif codigo_aei in indices_aei:
                        idx = contadores_aei.get(codigo_aei, 0)
                        contadores_aei[codigo_aei] = idx + 1
                        if idx not in indices_aei[codigo_aei]: fdel.append(i)
                elif m_oei and len(col0) > 10:
                    if not oei_activo: fdel.append(i)
                elif m_oei:
                    codigo_oei = m_oei.group()
                    oei_activo = codigo_oei in oei_list
                    if oei_activo:
                        if codigo_oei in indices_oei:
                            idx = contadores_oei.get(codigo_oei, 0)
                            contadores_oei[codigo_oei] = idx + 1
                            if idx not in indices_oei[codigo_oei]: fdel.append(i)
                    else: fdel.append(i)
            for i in reversed(fdel): table._element.remove(table.rows[i]._element)
            return len(fdel)

        e = 0
        for table in doc.tables:
            col_oei = detectar_col(table, r'OEI\.\d{2}')
            col_aei = detectar_col(table, r'AEI\.\d{2}\.\d{2}')
            if col_oei == 0 and col_aei == 0 and es_tabla_matriz(table): e += filtrar_tabla_matriz(table)
            elif col_oei >= 0 and col_aei >= 0 and col_oei != col_aei: e += filtrar_tabla_combinada(table, col_oei, col_aei)
            else:
                if col_oei >= 0: e += filtrar_tabla(table, r'OEI\.\d{2}', oei_list, indices_oei, col_oei)
                if col_aei >= 0: e += filtrar_tabla(table, r'AEI\.\d{2}\.\d{2}', aei_list, indices_aei, col_aei)
        print(f"  Filas eliminadas: {e}")

        def esc(c, t):
            for p in c.paragraphs:
                for r in p.runs: r.text = ''
            if c.paragraphs:
                if c.paragraphs[0].runs: c.paragraphs[0].runs[0].text = str(t)
                else: c.paragraphs[0].add_run(str(t))

        def fmt(v):
            if not v: return ''
            try:
                f = float(v)
                return str(int(f)) if f == int(f) else str(f)
            except: return str(v)

        mok = 0
        for table in doc.tables:
            cab = ' '.join(c.text for r in table.rows[:2] for c in r.cells)
            if 'Logros Esperados' not in cab: continue
            for row in table.rows:
                cs = row.cells
                if len(cs) < 6: continue
                m = re.match(r'^(OEI\.\d{2}|AEI\.\d{2}\.\d{2})$', cs[0].text.strip())
                if not m or m.group() not in metas_por_codigo: continue
                meta = metas_por_codigo[m.group()][0]
                esc(cs[3], fmt(next((meta[k] for k in ['año_base', 'anio_base'] if k in meta and meta[k]), '')))
                esc(cs[4], fmt(meta.get('valor_base', '')))
                for j, a in enumerate(anios):
                    if 5+j < len(cs): esc(cs[5+j], fmt(meta.get(f'meta_{a}', '')))
                mok += 1
        print(f"  Metas asignadas: {mok}")

        # PASO 3: Fichas en Orden Jerárquico
        print("\nPASO 3: Fichas (orden jerárquico)...")
        fichas_path = resolver_ruta(FRONTEND_DIR / 'fichas_tecnicas.json', PROJECT_ROOT / 'fichas_tecnicas.json')
        with open(fichas_path, 'r', encoding='utf-8') as f:
            fichas_db = json.load(f)
        
        p = doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)
        pt = doc.add_paragraph()
        rt = pt.add_run("ANEXO A - 6: FICHA TÉCNICA DE INDICADORES")
        rt.bold = True
        rt.font.size = Pt(12)

        plantilla_ficha = resolver_ruta(PLANTILLAS_DIR / 'Plantilla_de_ficha_técnica.docx', PROJECT_ROOT / 'Plantilla_de_ficha_técnica.docx')
        doc_plt = Document(str(plantilla_ficha))
        tbl_plt = doc_plt.tables[0]

        fg = 0
        for codigo in codigos_ordenados:
            if codigo not in fichas_db: continue
            ficha = fichas_db[codigo][0]
            if fg > 0: doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            doc.element.body.append(deepcopy(tbl_plt._element))
            tb = doc.tables[-1]
            
            def stxt(ri, ci, v):
                if ri < len(tb.rows) and ci < len(tb.rows[ri].cells):
                    cell = tb.rows[ri].cells[ci]
                    for p in cell.paragraphs:
                        for r in p.runs: r.text = ''
                    if cell.paragraphs:
                        if cell.paragraphs[0].runs: cell.paragraphs[0].runs[0].text = str(v)
                        else: cell.paragraphs[0].add_run(str(v))
            
            stxt(1, 1, ficha.get('objetivo_accion', ''))
            stxt(2, 1, ficha.get('nombre_indicador', ''))
            stxt(3, 1, ficha.get('justificacion', ''))
            stxt(4, 1, ficha.get('responsables', ''))
            stxt(5, 1, ficha.get('limitaciones', ''))
            stxt(6, 1, ficha.get('metodo_calculo', ''))
            stxt(7, 1, ficha.get('sentido_esperado', ''))
            stxt(8, 1, ficha.get('proceso_recoleccion', ''))
            stxt(9, 1, ficha.get('fuente_datos', ''))
            stxt(11, 1, ficha.get('anio_base', ''))
            stxt(12, 1, ficha.get('valor_relativo', ''))
            stxt(13, 1, ficha.get('valor_absoluto', ''))
            fg += 1

        print(f"\nTotal fichas generadas: {fg}")

        # --- Guardar documento final en un BytesIO (Memoria RAM) ---
        output_stream = io.BytesIO()
        doc.save(output_stream)
        output_bytes = output_stream.getvalue()

        output_name = f"PEI_{d['nombre_municipio'].replace(' ', '_') or 'Doc'}.docx"
        
        print(f"\n✅ GENERADO Y ENVIADO DESDE MEMORIA: {output_name}")
        print("="*70)
        
        return output_bytes, output_name


if __name__ == '__main__':
    cargar_catalogo_ues()
    
    # --- NUEVO: Configuración de Puerto para Render y Local ---
    # Render asigna automáticamente una variable de entorno llamada PORT.
    # Si no la encuentra (por ejemplo, en tu computadora local), usa el 8000.
    puerto = int(os.environ.get('PORT', 8000))
    
    # Para que Render pueda exponer tu aplicación a Internet, 
    # el host debe ser obligatoriamente '0.0.0.0' en lugar de '' o 'localhost'.
    host = '0.0.0.0'
    
    httpd = HTTPServer((host, puerto), partial(PEIHandler, directory=str(PROJECT_ROOT)))
    
    print("\n" + "="*70)
    print(f"  GENERADOR PEI - SERVIDOR ACTIVO")
    print(f"  Escuchando en http://{host}:{puerto}")
    print("="*70 + "\n")
    
    try:
        httpd.serve_forever()
    except KeyboardInterrupt:
        print("\nDetenido.")