#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from http.server import HTTPServer, SimpleHTTPRequestHandler
from urllib.parse import unquote, urlsplit
import json, math, os, re, sys, tempfile, unicodedata
from copy import deepcopy

from catalogo_ue import normalizar_id_ue, obtener_ue


WORD_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
MAX_PAYLOAD_BYTES = 1024 * 1024
MAX_MUNICIPALITY_LENGTH = 120
ORDENANZA_PLACEHOLDER = '«Ordenanza_PDC»'
ANEXO_FICHAS_MARKER = 'ANEXO A - 6'
PERIODO_PATTERN = re.compile(r'^\s*(\d{4})\s*[-–—]\s*(\d{4})\s*$')
STATIC_FILE_ALLOWLIST = {
    '/': 'index_fase4.html',
    '/index_fase4.html': 'index_fase4.html',
    '/matriz_estandar.json': 'matriz_estandar.json',
}


def _w_tag(name):
    return f'{{{WORD_NS}}}{name}'


def _local_name(element):
    return element.tag.rsplit('}', 1)[-1]


def _direct_child(element, name):
    return next((child for child in element if _local_name(child) == name), None)


def _direct_children(element, name):
    return [child for child in element if _local_name(child) == name]


def _xml_text(element):
    return ''.join(
        child.text or ''
        for child in element.iter()
        if _local_name(child) == 't'
    )


def _set_xml_text(element, value):
    text_nodes = [child for child in element.iter() if _local_name(child) == 't']
    if not text_nodes:
        return
    text_nodes[0].text = str(value)
    for text_node in text_nodes[1:]:
        text_node.text = ''


def _clear_xml_text(element):
    for child in element.iter():
        if _local_name(child) == 't':
            child.text = ''


def _grid_span(cell):
    tc_pr = _direct_child(cell, 'tcPr')
    if tc_pr is None:
        return 1
    grid_span = _direct_child(tc_pr, 'gridSpan')
    if grid_span is None:
        return 1
    try:
        return max(1, int(grid_span.get(_w_tag('val'), '1')))
    except (TypeError, ValueError):
        return 1


def _set_grid_span(cell, span):
    tc_pr = _direct_child(cell, 'tcPr')
    if tc_pr is None:
        return
    grid_span = _direct_child(tc_pr, 'gridSpan')
    if span <= 1:
        if grid_span is not None:
            tc_pr.remove(grid_span)
        return
    if grid_span is None:
        try:
            from docx.oxml import OxmlElement
            grid_span = OxmlElement('w:gridSpan')
        except ImportError:
            import xml.etree.ElementTree as ElementTree
            grid_span = ElementTree.Element(_w_tag('gridSpan'))
        tc_pr.append(grid_span)
    grid_span.set(_w_tag('val'), str(span))


def _row_column_count(row):
    return sum(_grid_span(cell) for cell in _direct_children(row, 'tc'))


def ajustar_columnas_anio(table, anios, header_row, first_year_col):
    """Resize a fixed-width Word table so every requested year has a column."""
    years = list(anios)
    if not years:
        raise ValueError('El período PEI no contiene años utilizables.')

    table_element = getattr(table, '_tbl', table)
    grid = _direct_child(table_element, 'tblGrid')
    if grid is None:
        raise ValueError('La tabla no contiene una cuadrícula de columnas.')

    grid_columns = _direct_children(grid, 'gridCol')
    desired_columns = first_year_col + len(years)
    if not grid_columns:
        raise ValueError('La tabla no contiene columnas configuradas.')

    if desired_columns > len(grid_columns):
        template_column = grid_columns[-1]
        for _ in range(desired_columns - len(grid_columns)):
            grid.append(deepcopy(template_column))
    elif desired_columns < len(grid_columns):
        for column in grid_columns[desired_columns:]:
            grid.remove(column)

    for row in _direct_children(table_element, 'tr'):
        logical_columns = _row_column_count(row)
        cells = _direct_children(row, 'tc')
        if not cells:
            continue

        if logical_columns < desired_columns:
            missing = desired_columns - logical_columns
            last_cell = cells[-1]
            last_span = _grid_span(last_cell)
            if last_span > 1:
                _set_grid_span(last_cell, last_span + missing)
            else:
                for _ in range(missing):
                    new_cell = deepcopy(last_cell)
                    _clear_xml_text(new_cell)
                    row.append(new_cell)
        elif logical_columns > desired_columns:
            excess = logical_columns - desired_columns
            while excess > 0:
                last_cell = _direct_children(row, 'tc')[-1]
                last_span = _grid_span(last_cell)
                if last_span > excess:
                    _set_grid_span(last_cell, last_span - excess)
                    excess = 0
                else:
                    row.remove(last_cell)
                    excess -= last_span

    header = _direct_children(table_element, 'tr')[header_row]
    header_cells = _direct_children(header, 'tc')
    for offset, year in enumerate(years, first_year_col):
        _set_xml_text(header_cells[offset], year)
    return {year: index for index, year in enumerate(years, first_year_col)}


def normalizar_bordes_tabla(table):
    """Ensure every target cell has a visible four-sided border."""
    table_element = getattr(table, '_tbl', table)
    for cell in (element for element in table_element.iter() if _local_name(element) == 'tc'):
        tc_pr = _direct_child(cell, 'tcPr')
        if tc_pr is None:
            continue
        borders = _direct_child(tc_pr, 'tcBorders')
        if borders is None:
            try:
                from docx.oxml import OxmlElement
                borders = OxmlElement('w:tcBorders')
            except ImportError:
                import xml.etree.ElementTree as ElementTree
                borders = ElementTree.Element(_w_tag('tcBorders'))
            tc_pr.append(borders)
        for side_name in ('top', 'left', 'bottom', 'right'):
            side = _direct_child(borders, side_name)
            if side is None:
                try:
                    from docx.oxml import OxmlElement
                    side = OxmlElement(f'w:{side_name}')
                except ImportError:
                    import xml.etree.ElementTree as ElementTree
                    side = ElementTree.Element(_w_tag(side_name))
                borders.append(side)
            if side.get(_w_tag('val')) in (None, 'nil'):
                side.set(_w_tag('val'), 'single')
            if side.get(_w_tag('sz')) is None:
                side.set(_w_tag('sz'), '4')
            if side.get(_w_tag('space')) is None:
                side.set(_w_tag('space'), '0')
            if side.get(_w_tag('color')) is None:
                side.set(_w_tag('color'), 'auto')


def extraer_nota_ficha(table):
    """Remove the template note row and return its original paragraph XML."""
    for row in _direct_children(table, 'tr'):
        if 'Nota:' not in _xml_text(row):
            continue
        cells = _direct_children(row, 'tc')
        paragraphs = [
            paragraph
            for paragraph in (_direct_children(cells[0], 'p') if cells else [])
            if _xml_text(paragraph).strip()
        ]
        note = deepcopy(paragraphs[0]) if paragraphs else None
        if note is not None:
            for paragraph in paragraphs[1:]:
                try:
                    from docx.oxml import OxmlElement
                    break_run = OxmlElement('w:r')
                    break_run.append(OxmlElement('w:br'))
                except ImportError:
                    import xml.etree.ElementTree as ElementTree
                    break_run = ElementTree.Element(_w_tag('r'))
                    break_run.append(ElementTree.Element(_w_tag('br')))
                note.append(break_run)
                for child in paragraph:
                    if _local_name(child) != 'pPr':
                        note.append(deepcopy(child))
        table.remove(row)
        return note
    return None


def insertar_despues_de_elemento(parent, reference, element):
    index = list(parent).index(reference)
    parent.insert(index + 1, element)


def eliminar_parrafos_con_texto(body, text):
    removed = 0
    for element in list(body):
        if _local_name(element) == 'p' and text in _xml_text(element):
            body.remove(element)
            removed += 1
    return removed


def buscar_parrafo_con_texto(body, text):
    return next(
        (element for element in body
         if _local_name(element) == 'p' and text in _xml_text(element)),
        None,
    )


def eliminar_parrafos_vacios_adyacentes(body, reference):
    index = list(body).index(reference)
    while index > 0 and _local_name(body[index - 1]) == 'p' and _parrafo_body_vacio_sin_salto(body[index - 1]):
        body.remove(body[index - 1])
        index -= 1
    while index + 1 < len(body) and _local_name(body[index + 1]) == 'p' and _parrafo_body_vacio_sin_salto(body[index + 1]):
        body.remove(body[index + 1])


def asegurar_salto_de_pagina_antes(paragraph):
    p_pr = _direct_child(paragraph, 'pPr')
    if p_pr is None:
        try:
            from docx.oxml import OxmlElement
            p_pr = OxmlElement('w:pPr')
        except ImportError:
            import xml.etree.ElementTree as ElementTree
            p_pr = ElementTree.Element(_w_tag('pPr'))
        paragraph.insert(0, p_pr)
    if _direct_child(p_pr, 'pageBreakBefore') is None:
        try:
            from docx.oxml import OxmlElement
            p_pr.append(OxmlElement('w:pageBreakBefore'))
        except ImportError:
            import xml.etree.ElementTree as ElementTree
            p_pr.append(ElementTree.Element(_w_tag('pageBreakBefore')))


def reemplazar_placeholder_docx(document, placeholder, value):
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for paragraph in paragraphs:
        if placeholder not in paragraph.text:
            continue
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value)
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, value)


def parsear_id_indicador(indicator_id):
    """Return (code, ordinal) from the ids emitted by the frontend."""
    patterns = (
        r'^ind-oei-(OEI\.\d+)-(\d+)$',
        r'^ind-aei-(AEI\.\d+\.\d+)-(\d+)$',
    )
    for pattern in patterns:
        match = re.fullmatch(pattern, indicator_id or '')
        if match:
            return match.group(1), int(match.group(2))
    return None


def indexar_metas(metas_raw):
    """Index metas by indicator code and the ordinal encoded in its id."""
    indexed = {}
    if not isinstance(metas_raw, dict):
        return indexed

    for indicator_id, meta in metas_raw.items():
        parsed = parsear_id_indicador(indicator_id)
        if parsed and isinstance(meta, dict):
            indexed[parsed] = meta
    return indexed


def _anio_de_ficha(ficha):
    value = ficha.get('anio_base', ficha.get('año_base', '')) if isinstance(ficha, dict) else ''
    match = re.search(r'\d{4}', str(value))
    return int(match.group()) if match else None


def _es_inicio_ficha(ficha):
    label = str(ficha.get('linea_base', '')).strip().lower() if isinstance(ficha, dict) else ''
    return label.startswith('línea') or label.startswith('linea')


def _es_nota_ficha(value):
    return isinstance(value, str) and value.strip().lower().startswith('nota:')


def valor_absoluto_ficha(ficha):
    """Return an absolute value, excluding source notes stored in that field."""
    value = ficha.get('valor_absoluto', '') if isinstance(ficha, dict) else ''
    return '' if _es_nota_ficha(value) else value


def obtener_nota_ficha(ficha):
    """Return the ficha note, including legacy notes stored as absolute values."""
    if not isinstance(ficha, dict):
        return ''
    note = ficha.get('nota', '')
    if _es_nota_ficha(note):
        return note
    absolute = ficha.get('valor_absoluto', '')
    return absolute if _es_nota_ficha(absolute) else ''


def indexar_fichas(fichas_db):
    """Index annual ficha records by (code, indicator ordinal, year).

    The source JSON stores one annual block per indicator but does not store
    the ordinal. Each block starts at its baseline record, so the ordinal is
    recovered from those deterministic block boundaries.
    """
    indexed = {}
    for codigo, records in (fichas_db or {}).items():
        ordinal = -1
        previous_year = None
        for ficha in records if isinstance(records, list) else []:
            year = _anio_de_ficha(ficha)
            if year is None:
                continue

            explicit_ordinal = ficha.get('ordinal_indicador') if isinstance(ficha, dict) else None
            if isinstance(explicit_ordinal, int):
                ordinal = explicit_ordinal
            elif ordinal < 0 or _es_inicio_ficha(ficha) or (
                previous_year is not None and year <= previous_year
            ):
                ordinal += 1

            indexed[(codigo, ordinal, year)] = ficha
            previous_year = year
    return indexed


def obtener_fichas_indicador(fichas_index, codigo, ordinal):
    return {
        year: ficha
        for (ficha_codigo, ficha_ordinal, year), ficha in fichas_index.items()
        if ficha_codigo == codigo and ficha_ordinal == ordinal
    }


def seleccionar_ficha_base(fichas_index, codigo, ordinal, anios_periodo):
    """Select the explicit baseline record, never the first PEI achievement year.

    The baseline column is separate from the period's achievement columns. The
    period may contain the same numeric year as the baseline, but that year is
    retained as an achievement instead of being silently replaced by it.
    """
    records = obtener_fichas_indicador(fichas_index, codigo, ordinal)
    if not records:
        return None
    baseline_records = {
        year: ficha
        for year, ficha in records.items()
        if _es_inicio_ficha(ficha)
    }
    if not baseline_records:
        return None
    return baseline_records[min(baseline_records)]


def extraer_anios_periodo(periodo):
    """Return every PEI year as an achievement year, including the start year.

    The fixed baseline column is not part of this list and must not cause a
    period year to be dropped or shifted. Missing source data remains blank in
    the corresponding achievement cell.
    """
    match = PERIODO_PATTERN.fullmatch(str(periodo or ''))
    if not match:
        return []
    start, end = (int(value) for value in match.groups())
    if start > end:
        raise ValueError('El período PEI tiene el año inicial posterior al año final.')
    return list(range(start, end + 1))


def mapear_columnas_anio(headers):
    """Map a year header to its real table column index."""
    columns = {}
    for index, header in enumerate(headers):
        match = re.search(r'\b(\d{4})\b', str(header or ''))
        if match:
            columns[int(match.group(1))] = index
    return columns


def resolver_anios_columnas(periodo, year_columns):
    """Return usable years and missing period years without inventing columns."""
    available = list(year_columns)
    requested = extraer_anios_periodo(periodo)
    if not requested:
        return available, []
    available_set = set(available)
    usable = [year for year in requested if year in available_set]
    missing = [year for year in requested if year not in available_set]
    return usable, missing


def construir_fichas_seleccionadas(codigos_ordenados, indices_oei, indices_aei, fichas_index):
    """Expand priority order into one ficha entry per selected indicator."""
    selected = []
    for codigo in codigos_ordenados:
        indices = indices_oei.get(codigo) if codigo.startswith('OEI.') else indices_aei.get(codigo)
        for ordinal in sorted(indices or ()):
            if any(
                ficha_codigo == codigo and ficha_ordinal == ordinal
                for ficha_codigo, ficha_ordinal, _ in fichas_index
            ):
                selected.append((codigo, ordinal))
    return selected


def _parrafo_body_vacio_sin_salto(paragraph):
    """True only for a body paragraph with no visible or structural content."""
    for element in paragraph.iter():
        local_name = element.tag.rsplit('}', 1)[-1]
        if local_name in {'p', 'pPr', 'r', 'rPr'}:
            continue
        if local_name == 't':
            if (element.text or '').strip():
                return False
            continue
        return False
    return True


def limpiar_parrafos_body_vacios_finales(body):
    """Remove trailing empty body paragraphs while preserving sectPr and breaks."""
    removed = 0
    for element in reversed(list(body)):
        local_name = element.tag.rsplit('}', 1)[-1]
        if local_name == 'sectPr':
            continue
        if local_name != 'p' or not _parrafo_body_vacio_sin_salto(element):
            break
        body.remove(element)
        removed += 1
    return removed


def insertar_antes_de_sectpr(body, element):
    for index, child in enumerate(body):
        if child.tag == _w_tag('sectPr'):
            body.insert(index, element)
            return
    body.append(element)


class PayloadValidationError(ValueError):
    """Validation errors that can be returned to the JSON client."""

    def __init__(self, errors):
        self.errors = errors
        super().__init__('El payload contiene datos inválidos')


def cargar_matriz_estandar():
    path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'matriz_estandar.json')
    with open(path, 'r', encoding='utf-8') as matriz_file:
        return json.load(matriz_file)


def sanitizar_nombre_municipio(nombre):
    """Return a useful filename fragment without filesystem control syntax."""
    if not isinstance(nombre, str):
        raise ValueError('El nombre del municipio debe ser texto.')

    normalized = unicodedata.normalize('NFC', nombre).strip()
    if not normalized:
        raise ValueError('El nombre del municipio es obligatorio.')
    if len(normalized) > MAX_MUNICIPALITY_LENGTH:
        raise ValueError('El nombre del municipio es demasiado extenso.')
    if any(char in normalized for char in ('/', '\\')) or '..' in normalized:
        raise ValueError('El nombre del municipio contiene una ruta inválida.')
    if any(unicodedata.category(char).startswith('C') for char in normalized):
        raise ValueError('El nombre del municipio contiene caracteres inválidos.')

    safe = ''.join(
        char if char.isalnum() or char in '._-()' else '_'
        for char in normalized
    )
    safe = re.sub(r'_+', '_', re.sub(r'\s+', '_', safe)).strip('._-')
    if not safe:
        raise ValueError('El nombre del municipio no permite crear un archivo.')
    return safe


def _registrar_codigo(values, raw_value, prefix, field, errors):
    if not isinstance(raw_value, str) or not raw_value.startswith(prefix):
        errors.append({
            'field': field,
            'code': 'invalid_format',
            'message': f'{field} contiene un identificador inválido: {raw_value!r}'
        })
        return None

    value = raw_value[len(prefix):]
    if value in values:
        errors.append({
            'field': field,
            'code': 'duplicate',
            'message': f'{field} contiene un código repetido: {value}'
        })
        return None

    values.add(value)
    return value


def _validar_metas(metas, errors):
    if not isinstance(metas, dict):
        errors.append({
            'field': 'metas',
            'code': 'invalid_type',
            'message': 'Las metas deben ser un objeto indexado por indicador.'
        })
        return

    for indicator_id, meta in metas.items():
        field = f'metas.{indicator_id}'
        if not isinstance(indicator_id, str) or not indicator_id:
            errors.append({
                'field': 'metas',
                'code': 'invalid_type',
                'message': 'Cada clave de metas debe ser un identificador de indicador.'
            })
            continue
        if not isinstance(meta, dict):
            errors.append({
                'field': field,
                'code': 'invalid_type',
                'message': 'Los datos de cada indicador deben ser un objeto.'
            })
            continue

        for name, value in meta.items():
            if value is None or value == '':
                continue
            if isinstance(value, bool):
                valid = False
            elif isinstance(value, (int, float)):
                valid = math.isfinite(value)
            elif isinstance(value, str):
                try:
                    valid = math.isfinite(float(value.strip()))
                except ValueError:
                    valid = False
            else:
                valid = False
            if not valid:
                errors.append({
                    'field': f'{field}.{name}',
                    'code': 'invalid_type',
                    'message': 'El valor de la meta debe ser numérico o estar vacío.'
                })


def validar_payload(data, matriz=None):
    """Validate selection hierarchy and return normalized priority order."""
    if not isinstance(data, dict):
        raise PayloadValidationError([{
            'field': 'payload',
            'code': 'invalid_type',
            'message': 'El cuerpo debe ser un objeto JSON.'
        }])

    matriz = matriz or cargar_matriz_estandar()
    oei_by_code = {}
    aei_by_code = {}
    indicador_oei_count = {}
    indicador_aei_count = {}

    for oei in matriz.get('oei', []):
        codigo_oei = oei.get('codigo')
        oei_by_code[codigo_oei] = oei
        indicador_oei_count[codigo_oei] = len(oei.get('indicadores', []))
        for aei in oei.get('aei', []):
            codigo_aei = aei.get('codigo')
            aei_by_code[codigo_aei] = {
                'oei': codigo_oei,
                'indicadores': len(aei.get('indicadores', []))
            }
            indicador_aei_count[codigo_aei] = len(aei.get('indicadores', []))

    errors = []
    if 'codigo_ue' in data:
        codigo_ue = data.get('codigo_ue')
        if not isinstance(codigo_ue, str):
            errors.append({
                'field': 'codigo_ue',
                'code': 'invalid_type',
                'message': 'El código UE debe ser texto.'
            })
        elif not codigo_ue.strip():
            errors.append({
                'field': 'codigo_ue',
                'code': 'required',
                'message': 'El código UE es obligatorio.'
            })
        else:
            try:
                normalizar_id_ue(codigo_ue.strip())
            except ValueError as error:
                errors.append({
                    'field': 'codigo_ue',
                    'code': 'invalid_format',
                    'message': str(error)
                })

    if 'metas' in data:
        _validar_metas(data.get('metas'), errors)

    periodo_value = data.get('periodo_pei', '')
    if not isinstance(periodo_value, str):
        errors.append({
            'field': 'periodo_pei',
            'code': 'invalid_type',
            'message': 'El período PEI debe ser texto.',
        })
        periodo = ''
    else:
        periodo = periodo_value.strip()
    if isinstance(periodo_value, str) and not periodo:
        errors.append({
            'field': 'periodo_pei',
            'code': 'required',
            'message': 'El período PEI es obligatorio.',
        })
    else:
        try:
            if not extraer_anios_periodo(periodo):
                raise ValueError('El período PEI debe tener el formato AAAA-AAAA.')
        except ValueError as error:
            errors.append({
                'field': 'periodo_pei',
                'code': 'invalid_period',
                'message': str(error),
            })

    ordenanza_value = data.get('ordenanza_pdc', '')
    ordenanza = ordenanza_value.strip() if isinstance(ordenanza_value, str) else ''
    if not ordenanza:
        errors.append({
            'field': 'ordenanza_pdc',
            'code': 'required',
            'message': 'El número de ordenanza es obligatorio.',
        })
    elif 'ordenanza_pdc' in ordenanza.lower() or ORDENANZA_PLACEHOLDER in ordenanza:
        errors.append({
            'field': 'ordenanza_pdc',
            'code': 'placeholder',
            'message': 'El número de ordenanza no puede contener un placeholder.',
        })
    selecciones = data.get('selecciones')
    prioridades = data.get('prioridades')
    if not isinstance(selecciones, dict):
        errors.append({
            'field': 'selecciones',
            'code': 'required',
            'message': 'Se requiere el objeto selecciones.'
        })
        selecciones = {}
    if not isinstance(prioridades, dict):
        errors.append({
            'field': 'prioridades',
            'code': 'required',
            'message': 'Se requiere el objeto prioridades.'
        })
        prioridades = {}

    selected_oei = []
    selected_oei_set = set()
    selected_aei = []
    selected_aei_set = set()

    raw_oei = selecciones.get('oei', [])
    raw_aei = selecciones.get('aei', [])
    if not isinstance(raw_oei, list):
        errors.append({'field': 'selecciones.oei', 'code': 'invalid_type', 'message': 'Debe ser una lista.'})
        raw_oei = []
    if not isinstance(raw_aei, list):
        errors.append({'field': 'selecciones.aei', 'code': 'invalid_type', 'message': 'Debe ser una lista.'})
        raw_aei = []

    for raw_code in raw_oei:
        code = _registrar_codigo(selected_oei_set, raw_code, 'oei-', 'selecciones.oei', errors)
        if code is None:
            continue
        if code not in oei_by_code:
            errors.append({
                'field': 'selecciones.oei',
                'code': 'unknown_code',
                'message': f'El OEI no existe en la matriz: {code}'
            })
            continue
        selected_oei.append(code)

    for raw_code in raw_aei:
        code = _registrar_codigo(selected_aei_set, raw_code, 'aei-', 'selecciones.aei', errors)
        if code is None:
            continue
        relation = aei_by_code.get(code)
        if relation is None:
            errors.append({
                'field': 'selecciones.aei',
                'code': 'unknown_code',
                'message': f'El AEI no existe en la matriz: {code}'
            })
            continue
        if relation['oei'] not in selected_oei_set:
            errors.append({
                'field': 'selecciones.aei',
                'code': 'missing_parent',
                'message': f'El AEI {code} requiere seleccionar su OEI padre {relation["oei"]}.'
            })
            continue
        selected_aei.append(code)

    raw_ind_oei = selecciones.get('indicadoresOEI', [])
    raw_ind_aei = selecciones.get('indicadoresAEI', [])
    if not isinstance(raw_ind_oei, list):
        errors.append({'field': 'selecciones.indicadoresOEI', 'code': 'invalid_type', 'message': 'Debe ser una lista.'})
        raw_ind_oei = []
    if not isinstance(raw_ind_aei, list):
        errors.append({'field': 'selecciones.indicadoresAEI', 'code': 'invalid_type', 'message': 'Debe ser una lista.'})
        raw_ind_aei = []

    indicator_patterns = (
        ('selecciones.indicadoresOEI', raw_ind_oei, re.compile(r'^ind-oei-(OEI\.\d+)-(\d+)$'), selected_oei_set),
        ('selecciones.indicadoresAEI', raw_ind_aei, re.compile(r'^ind-aei-(AEI\.\d+\.\d+)-(\d+)$'), selected_aei_set),
    )
    seen_indicators = set()
    for field, raw_indicators, pattern, selected_parents in indicator_patterns:
        for raw_indicator in raw_indicators:
            match = pattern.fullmatch(raw_indicator) if isinstance(raw_indicator, str) else None
            if not match:
                errors.append({
                    'field': field,
                    'code': 'invalid_format',
                    'message': f'{field} contiene un identificador inválido: {raw_indicator!r}'
                })
                continue

            parent_code, raw_index = match.groups()
            try:
                index = int(raw_index)
            except ValueError:
                index = -1
            if parent_code not in selected_parents:
                errors.append({
                    'field': field,
                    'code': 'missing_parent',
                    'message': f'El indicador {raw_indicator} requiere seleccionar su padre.'
                })
                continue

            max_index = (indicador_oei_count if field.endswith('OEI') else indicador_aei_count).get(parent_code)
            if max_index is None:
                errors.append({
                    'field': field,
                    'code': 'unknown_code',
                    'message': f'El código padre del indicador no existe: {parent_code}'
                })
            elif index < 0 or index >= max_index:
                errors.append({
                    'field': field,
                    'code': 'unknown_code',
                    'message': f'El índice del indicador no existe: {raw_indicator}'
                })

            if raw_indicator in seen_indicators:
                errors.append({
                    'field': field,
                    'code': 'duplicate',
                    'message': f'Indicador repetido: {raw_indicator}'
                })
            seen_indicators.add(raw_indicator)

    raw_oei_priority = prioridades.get('oei')
    raw_aei_priority = prioridades.get('aei')
    if not isinstance(raw_oei_priority, list):
        errors.append({'field': 'prioridades.oei', 'code': 'required', 'message': 'Debe ser una lista ordenada de códigos OEI.'})
        raw_oei_priority = []
    if not isinstance(raw_aei_priority, dict):
        errors.append({'field': 'prioridades.aei', 'code': 'required', 'message': 'Debe ser un objeto ordenado por OEI.'})
        raw_aei_priority = {}

    priority_oei = []
    priority_oei_set = set()
    for code in raw_oei_priority:
        if not isinstance(code, str) or code not in oei_by_code:
            errors.append({
                'field': 'prioridades.oei',
                'code': 'unknown_code',
                'message': f'El código de prioridad OEI no existe: {code!r}'
            })
            continue
        if code in priority_oei_set:
            errors.append({'field': 'prioridades.oei', 'code': 'duplicate', 'message': f'OEI repetido en prioridades: {code}'})
            continue
        priority_oei_set.add(code)
        priority_oei.append(code)

    if priority_oei_set != selected_oei_set:
        errors.append({
            'field': 'prioridades.oei',
            'code': 'selection_mismatch',
            'message': 'prioridades.oei debe contener exactamente los OEI seleccionados.'
        })

    priority_aei = {}
    for codigo_oei, raw_codes in raw_aei_priority.items():
        if codigo_oei not in selected_oei_set:
            errors.append({
                'field': 'prioridades.aei',
                'code': 'invalid_parent',
                'message': f'La prioridad AEI referencia un OEI no seleccionado: {codigo_oei}'
            })
            continue
        if not isinstance(raw_codes, list):
            errors.append({
                'field': f'prioridades.aei.{codigo_oei}',
                'code': 'invalid_type',
                'message': 'Debe ser una lista ordenada de códigos AEI.'
            })
            continue

        seen_codes = set()
        normalized_codes = []
        for codigo_aei in raw_codes:
            relation = aei_by_code.get(codigo_aei) if isinstance(codigo_aei, str) else None
            if relation is None:
                errors.append({
                    'field': f'prioridades.aei.{codigo_oei}',
                    'code': 'unknown_code',
                    'message': f'El AEI no existe en la matriz: {codigo_aei}'
                })
                continue
            if relation['oei'] != codigo_oei:
                errors.append({
                    'field': f'prioridades.aei.{codigo_oei}',
                    'code': 'invalid_relation',
                    'message': f'El AEI {codigo_aei} no pertenece al OEI {codigo_oei}.'
                })
                continue
            if codigo_aei not in selected_aei_set:
                errors.append({
                    'field': f'prioridades.aei.{codigo_oei}',
                    'code': 'selection_mismatch',
                    'message': f'El AEI {codigo_aei} no está seleccionado.'
                })
                continue
            if codigo_aei in seen_codes:
                errors.append({
                    'field': f'prioridades.aei.{codigo_oei}',
                    'code': 'duplicate',
                    'message': f'AEI repetido en prioridades: {codigo_aei}'
                })
                continue
            seen_codes.add(codigo_aei)
            normalized_codes.append(codigo_aei)
        priority_aei[codigo_oei] = normalized_codes

    for codigo_oei in selected_oei:
        if codigo_oei not in raw_aei_priority:
            errors.append({
                'field': f'prioridades.aei.{codigo_oei}',
                'code': 'required',
                'message': f'Falta el orden AEI del OEI seleccionado {codigo_oei}.'
            })
            priority_aei[codigo_oei] = []
        expected = {codigo for codigo in selected_aei if aei_by_code[codigo]['oei'] == codigo_oei}
        actual = set(priority_aei.get(codigo_oei, []))
        if expected != actual:
            errors.append({
                'field': f'prioridades.aei.{codigo_oei}',
                'code': 'selection_mismatch',
                'message': f'El orden AEI de {codigo_oei} debe contener exactamente sus AEI seleccionadas.'
            })

    if errors:
        raise PayloadValidationError(errors)

    return {'oei': priority_oei, 'aei': priority_aei}


class PEIHandler(SimpleHTTPRequestHandler):
    def do_GET(self):
        path = urlsplit(self.path).path
        if path == '/api/ue' or path.startswith('/api/ue/'):
            self._get_unidad_ejecutora(path)
            return

        allowed_file = self._allowed_static_file(path)
        if allowed_file is None:
            self._send_json(404, {
                'success': False,
                'error': {
                    'code': 'not_found',
                    'message': 'Recurso no encontrado.',
                },
            })
            return

        original_path = self.path
        self.path = '/' + allowed_file
        try:
            super().do_GET()
        finally:
            self.path = original_path

    def _allowed_static_file(self, path):
        allowed_file = STATIC_FILE_ALLOWLIST.get(path)
        if allowed_file is not None:
            return allowed_file

        decoded_path = unquote(path)
        if not decoded_path.startswith('/PEI_') or not decoded_path.endswith('.docx'):
            return None

        filename = decoded_path[1:]
        if '/' in filename:
            return None

        try:
            safe_name = sanitizar_nombre_municipio(filename[4:-5])
        except ValueError:
            return None

        if filename != f'PEI_{safe_name}.docx' or not os.path.isfile(filename):
            return None
        return filename

    def _get_unidad_ejecutora(self, path):
        parts = path.split('/')
        if len(parts) != 4 or parts[1:3] != ['api', 'ue'] or not parts[3]:
            self._send_json(400, {
                'success': False,
                'error': {
                    'code': 'invalid_id',
                    'message': 'La ruta debe tener el formato /api/ue/<id>.',
                },
            })
            return

        try:
            id_ue = normalizar_id_ue(parts[3])
        except ValueError as error:
            self._send_json(400, {
                'success': False,
                'error': {'code': 'invalid_id', 'message': str(error)},
            })
            return

        try:
            data = obtener_ue(id_ue)
        except Exception:
            self._send_json(500, {
                'success': False,
                'error': {
                    'code': 'internal_error',
                    'message': 'Error interno del servidor.',
                },
            })
            return

        if data is None:
            self._send_json(404, {
                'success': False,
                'error': {
                    'code': 'ue_not_found',
                    'message': 'La UE no existe en el catalogo.',
                },
            })
            return

        self._send_json(200, {'success': True, 'data': data})

    def do_POST(self):
        if self.path != '/generar':
            self._send_json(404, {
                'success': False,
                'error': {
                    'code': 'not_found',
                    'message': 'Ruta no encontrada.',
                },
            })
            return

        content_length_header = self.headers.get('Content-Length')
        try:
            if content_length_header is None:
                raise ValueError('Falta el encabezado Content-Length.')
            content_length = int(content_length_header)
            if content_length < 0:
                raise ValueError('Content-Length inválido.')
        except ValueError as error:
            self._send_json(400, {
                'success': False,
                'error': 'Solicitud inválida.',
                'details': [{'field': 'Content-Length', 'code': 'invalid_length', 'message': str(error)}]
            })
            return

        if content_length > MAX_PAYLOAD_BYTES:
            self._send_json(413, {
                'success': False,
                'error': 'El payload excede el límite permitido.',
                'details': [{'field': 'payload', 'code': 'payload_too_large'}]
            })
            return

        if self.headers.get_content_type() != 'application/json':
            self._send_json(415, {
                'success': False,
                'error': 'El Content-Type debe ser application/json.',
                'details': [{'field': 'Content-Type', 'code': 'unsupported_media_type'}]
            })
            return

        try:
            raw_body = self.rfile.read(content_length)
            if len(raw_body) != content_length:
                raise ValueError('El cuerpo no coincide con Content-Length.')
            data = json.loads(raw_body.decode('utf-8'))
        except (ValueError, json.JSONDecodeError) as error:
            self._send_json(400, {
                'success': False,
                'error': 'Solicitud JSON inválida.',
                'details': [{'field': 'payload', 'code': 'invalid_json', 'message': str(error)}]
            })
            return

        try:
            validar_payload(data)
            output = self.generar_documento(data)
            self._send_json(200, {'success': True, 'file': output, 'message': 'Generado'})
        except PayloadValidationError as error:
            self._send_json(422, {
                'success': False,
                'error': str(error),
                'details': error.errors
            })
        except Exception as error:
            print(f'ERROR: {error}', file=sys.stderr)
            self._send_json(500, {
                'success': False,
                'error': 'Error interno del servidor.',
            })

    def _send_json(self, status, payload):
        body = json.dumps(payload, ensure_ascii=False).encode('utf-8')
        self.send_response(status)
        self.send_header('Content-type', 'application/json; charset=utf-8')
        self.send_header('Content-Length', str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def log_message(self, format, *args):
        pass

    def generar_documento(self, data, *, template_dir=None, output_path=None):
        fd, temp1 = tempfile.mkstemp(prefix='_pei_', suffix='.docx')
        os.close(fd)
        try:
            options = {}
            if template_dir is not None:
                options['template_dir'] = template_dir
            if output_path is not None:
                options['output_path'] = output_path
            return self._generar_documento(data, temp1, **options)
        finally:
            if os.path.exists(temp1):
                os.remove(temp1)

    def _generar_documento(self, data, temp1, *, template_dir=None, output_path=None):
        print("\n" + "="*70)
        print("GENERANDO DOCUMENTO PEI")
        print("="*70)

        prioridades = validar_payload(data)

        from mailmerge import MailMerge
        from docx import Document
        from docx.shared import Pt
        from copy import deepcopy

        template_dir = template_dir or os.getcwd()
        informe_template_path = os.path.join(template_dir, 'PEI_Estandar_-_Informe.docx')
        ficha_template_path = os.path.join(template_dir, 'Plantilla_de_ficha_técnica.docx')

        # Datos
        campos = ['codigo_ue', 'nombre_municipio', 'periodo_pei', 'nombre_provincia',
                  'nombre_region', 'nombre_alcalde', 'resolucion_alcaldia',
                  'plan_desarrollo_concertado', 'ordenanza_pdc', 'mision_institucional',
                  'politica_general_gobierno', 'decreto_politica_general_gobierno']
        d = {k: str(data.get(k, '')).strip() for k in campos}
        print(f"Municipio: {d['nombre_municipio']}")

        # OEI/AEI - Mantener orden de selección
        sel = data.get('selecciones', {})
        oei_list = [x.replace('oei-', '') for x in sel.get('oei', [])]
        aei_list = [x.replace('aei-', '') for x in sel.get('aei', [])]
        
        # Obtener prioridades en el orden enviado por el usuario.
        oei_ordenados = prioridades['oei']
        aei_ordenados = prioridades['aei']
        
        print(f"OEI ordenados: {oei_ordenados}")
        print(f"AEI totales: {len(aei_list)}")

        # Crear orden jerárquico: OEI → sus AEI
        codigos_ordenados = []
        for oei_code in oei_ordenados:
            # Agregar el OEI
            codigos_ordenados.append(oei_code)
            
            # Agregar sus AEI (las que empiecen con el código del OEI)
            # Ej: OEI.01 → AEI.01.01, AEI.01.02, etc.
            oei_num = oei_code.split('.')[1]  # '01' de 'OEI.01'
            aei_de_este_oei = [
                aei for aei in aei_list 
                if aei.startswith(f'AEI.{oei_num}.')
            ]
            aei_de_este_oei = [codigo for codigo in aei_ordenados[oei_code] if codigo in aei_de_este_oei]
            codigos_ordenados.extend(aei_de_este_oei)

        oei_prioridad_numero = {codigo: index + 1 for index, codigo in enumerate(oei_ordenados)}
        aei_prioridad_numero = {
            codigo_oei: {codigo: index + 1 for index, codigo in enumerate(codigos_aei)}
            for codigo_oei, codigos_aei in aei_ordenados.items()
        }
        
        print(f"\nOrden de fichas: {codigos_ordenados}")

        # Metas
        metas_raw = data.get('metas', {})
        metas_por_indicador = indexar_metas(metas_raw)
        anios_periodo = extraer_anios_periodo(d['periodo_pei'])

        # PASO 1
        print("\nPASO 1: MailMerge...")
        merge_data = {
            'Codigo_UE': d['codigo_ue'], 'Nombre_Municipio': d['nombre_municipio'],
            'Periodo_PEI': d['periodo_pei'], 'Nombre_Provincia': d['nombre_provincia'],
            'Nombre_Region': d['nombre_region'], 'Nombre_Alcalde': d['nombre_alcalde'],
            'Resolucion_Alcaldia': d['resolucion_alcaldia'],
            'Plan_Desarrollo_Concertado': d['plan_desarrollo_concertado'],
            'Ordenanza_PDC': d['ordenanza_pdc'], 'Mision_Institucional': d['mision_institucional'],
            'Politica_General_Gobierno': d['politica_general_gobierno'],
            'Decreto_Politica_General_Gobierno': d['decreto_politica_general_gobierno'],
        }
        with MailMerge(informe_template_path) as doc:
            doc.merge(**merge_data)
            doc.write(temp1)
        print("  OK")

        # PASO 2
        print("\nPASO 2: Filtrar...")
        doc = Document(temp1)
        reemplazar_placeholder_docx(doc, ORDENANZA_PLACEHOLDER, d['ordenanza_pdc'])

        # Construir mapa de indices seleccionados por codigo
        ind_oei_sel = sel.get('indicadoresOEI', [])
        ind_aei_sel = sel.get('indicadoresAEI', [])

        indices_oei = {}
        for ind_id in ind_oei_sel:
            m = re.match(r'ind-oei-(OEI\.\d+)-(\d+)', ind_id)
            if m:
                codigo, idx = m.group(1), int(m.group(2))
                indices_oei.setdefault(codigo, set()).add(idx)

        indices_aei = {}
        for ind_id in ind_aei_sel:
            m = re.match(r'ind-aei-(AEI\.\d+\.\d+)-(\d+)', ind_id)
            if m:
                codigo, idx = m.group(1), int(m.group(2))
                indices_aei.setdefault(codigo, set()).add(idx)

        aei_seleccionadas_por_oei = {}
        for codigo_aei in aei_list:
            match = re.match(r'^AEI\.(\d{2})\.\d{2}$', codigo_aei)
            if match:
                codigo_oei = f'OEI.{match.group(1)}'
                aei_seleccionadas_por_oei.setdefault(codigo_oei, set()).add(codigo_aei)

        indicador_ordinal_por_fila = {}
        indicador_ordinal_por_clave = {}

        def registrar_ordinal(row, codigo, ordinal):
            indicador_ordinal_por_fila[id(row._tr)] = ordinal
            indicador_ordinal_por_clave[
                (codigo, tuple(cell.text.strip() for cell in row.cells))
            ] = ordinal

        print(f'  Indices OEI: {indices_oei}')
        print(f'  Indices AEI: {indices_aei}')

        def filtrar_tabla(table, patron, lista, indices_map, col_codigo):
            """Filtra filas de una tabla buscando el codigo en una columna especifica."""
            fdel = []
            contadores = {}
            for i, row in enumerate(table.rows):
                if len(row.cells) <= col_codigo:
                    continue
                txt_col = row.cells[col_codigo].text.strip()
                m = re.search(patron, txt_col)
                if m:
                    if txt_col.lower().startswith('acciones estratégicas'):
                        codigo_oei = m.group()
                        if not aei_seleccionadas_por_oei.get(codigo_oei):
                            fdel.append(i)
                        continue
                    codigo = m.group()
                    idx = contadores.get(codigo, 0)
                    contadores[codigo] = idx + 1
                    registrar_ordinal(row, codigo, idx)
                    if codigo not in lista or idx not in indices_map.get(codigo, set()):
                        fdel.append(i)
            for i in reversed(fdel):
                table._element.remove(table.rows[i]._element)
            return len(fdel)

        def detectar_col(table, patron):
            """Detecta en que columna aparece el patron en una tabla."""
            for row in table.rows[1:4]:
                for c_idx, cell in enumerate(row.cells):
                    if re.search(patron, cell.text.strip()):
                        return c_idx
            return -1

        def filtrar_tabla_combinada(table, col_oei, col_aei):
            """Para tablas como Ruta Estrategica donde cada fila tiene OEI y AEI.
            Filtra, reordena y actualiza los números literales de prioridad."""
            filas_validas = []
            primera_fila_datos = None
            oei_actual = None
            for i, row in enumerate(table.rows):
                ncols = len(row.cells)
                if ncols <= max(col_oei, col_aei):
                    continue
                txt_oei = row.cells[col_oei].text.strip()
                txt_aei = row.cells[col_aei].text.strip()
                m_oei = re.search(r'OEI\.\d{2}', txt_oei)
                m_aei = re.search(r'AEI\.\d{2}\.\d{2}', txt_aei)
                if m_oei or m_aei:
                    if m_oei:
                        oei_actual = m_oei.group()
                    codigo_oei = m_oei.group() if m_oei else oei_actual
                    codigo_aei = m_aei.group() if m_aei else None
                    oei_ok = codigo_oei in oei_list if codigo_oei else False
                    aei_ok = (not m_aei) or (m_aei.group() in aei_list)
                    if oei_ok and aei_ok:
                        if primera_fila_datos is None:
                            primera_fila_datos = i
                        if codigo_oei:
                            prioridad_oei = oei_prioridad_numero[codigo_oei]
                            if col_oei > 0 and row.cells[col_oei - 1].text.strip().isdigit():
                                row.cells[col_oei - 1].text = str(prioridad_oei)
                        else:
                            prioridad_oei = 0
                        if codigo_aei:
                            prioridad_aei = aei_prioridad_numero[codigo_oei][codigo_aei]
                            if col_aei > 0 and row.cells[col_aei - 1].text.strip().isdigit():
                                row.cells[col_aei - 1].text = str(prioridad_aei)
                        else:
                            prioridad_aei = 0
                        filas_validas.append((
                            prioridad_oei,
                            prioridad_aei,
                            row._tr
                        ))

            filas_validas = sorted(filas_validas, key=lambda item: (item[0], item[1]))
            todas_las_filas = list(table.rows)
            filas_a_reubicar = {row_element for _, _, row_element in filas_validas}
            for row in reversed(todas_las_filas):
                if row._tr in filas_a_reubicar:
                    table._element.remove(row._tr)
            if primera_fila_datos is not None:
                filas_actuales = list(table.rows)
                ancla = filas_actuales[min(primera_fila_datos, len(filas_actuales) - 1)]._tr
                for _, _, row_element in filas_validas:
                    ancla.addnext(row_element)
                    ancla = row_element

            filas_validas_codes = {row_element for _, _, row_element in filas_validas}
            eliminadas = 0
            for row in list(table.rows):
                txt = ' '.join(cell.text for cell in row.cells)
                if re.search(r'OEI\.\d{2}|AEI\.\d{2}\.\d{2}', txt) and row._tr not in filas_validas_codes:
                    table._element.remove(row._tr)
                    eliminadas += 1
            return eliminadas

        def es_tabla_matriz(table):
            """Detecta si es la tabla B-3 Matriz PEI: tiene OEI y AEI en col 0
            y filas separadoras de Acciones Estrategicas."""
            for row in table.rows[2:5]:
                if row.cells and re.search(r'OEI\.\d{2}', row.cells[0].text):
                    return True
            return False

        def filtrar_tabla_matriz(table):
            """Filtra la tabla B-3: elimina grupos OEI+separador+AEIs no seleccionados."""
            fdel = []
            oei_activo = False
            contadores_oei = {}
            contadores_aei = {}

            for i, row in enumerate(table.rows):
                if not row.cells:
                    continue
                col0 = row.cells[0].text.strip()

                m_aei = re.search(r'AEI\.\d{2}\.\d{2}', col0)
                m_oei = re.search(r'OEI\.\d{2}', col0)

                if col0.lower().startswith('acciones estratégicas'):
                    codigo_oei = m_oei.group() if m_oei else None
                    if not codigo_oei or not aei_seleccionadas_por_oei.get(codigo_oei):
                        fdel.append(i)
                    continue

                # Fila AEI pura
                if m_aei:
                    codigo_aei = m_aei.group()
                    if not oei_activo or codigo_aei not in aei_list:
                        fdel.append(i)
                    else:
                        idx = contadores_aei.get(codigo_aei, 0)
                        contadores_aei[codigo_aei] = idx + 1
                        registrar_ordinal(row, codigo_aei, idx)
                        if idx not in indices_aei.get(codigo_aei, set()):
                            fdel.append(i)

                # Fila separadora "Acciones Estrategicas Institucionales del OEI.XX"
                elif m_oei and len(col0) > 10:
                    if not oei_activo:
                        fdel.append(i)

                # Fila OEI pura (codigo corto como "OEI.02")
                elif m_oei:
                    codigo_oei = m_oei.group()
                    oei_activo = codigo_oei in oei_list
                    if oei_activo:
                        idx = contadores_oei.get(codigo_oei, 0)
                        contadores_oei[codigo_oei] = idx + 1
                        registrar_ordinal(row, codigo_oei, idx)
                        if idx not in indices_oei.get(codigo_oei, set()):
                            fdel.append(i)
                    else:
                        fdel.append(i)

            for i in reversed(fdel):
                table._element.remove(table.rows[i]._element)
            return len(fdel)

        e = 0
        for table in doc.tables:
            cab = ' '.join(c.text for r in table.rows[:2] for c in r.cells)
            if 'Logros Esperados' in cab:
                # OEI and AEI rows share column zero in this table. Process
                # both patterns independently so each keeps its ordinal.
                e += filtrar_tabla(table, r'OEI\.\d{2}', oei_list, indices_oei, 0)
                e += filtrar_tabla(table, r'AEI\.\d{2}\.\d{2}', aei_list, indices_aei, 0)
                continue
            col_oei = detectar_col(table, r'OEI\.\d{2}')
            col_aei = detectar_col(table, r'AEI\.\d{2}\.\d{2}')

            if col_oei == 0 and col_aei == 0 and es_tabla_matriz(table):
                # Tabla B-3: OEI y AEI ambos en col 0, con separadores
                e += filtrar_tabla_matriz(table)
            elif col_oei >= 0 and col_aei >= 0 and col_oei != col_aei:
                # Tabla combinada: OEI y AEI en distintas columnas (Ruta Estrategica)
                e += filtrar_tabla_combinada(table, col_oei, col_aei)
            else:
                # Tabla simple: solo OEI o solo AEI
                if col_oei >= 0:
                    e += filtrar_tabla(table, r'OEI\.\d{2}', oei_list, indices_oei, col_oei)
                if col_aei >= 0:
                    e += filtrar_tabla(table, r'AEI\.\d{2}\.\d{2}', aei_list, indices_aei, col_aei)
        if doc.tables:
            normalizar_bordes_tabla(doc.tables[0])
        print(f"  Filas: {e}")

        # Metas
        def esc(c, t):
            for p in c.paragraphs:
                for r in p.runs: r.text = ''
            if c.paragraphs:
                runs = c.paragraphs[0].runs
                if runs: runs[0].text = str(t)
                else: c.paragraphs[0].add_run(str(t))

        def fmt(v):
            if v is None or v == '': return ''
            try:
                f = float(v)
                return str(int(f)) if f == int(f) else str(f)
            except: return str(v)

        def get(m, ks):
            for k in ks:
                if k in m and m[k] is not None: return m[k]
            return ''

        mok = 0
        for table in doc.tables:
            cab = ' '.join(c.text for r in table.rows[:2] for c in r.cells)
            if 'Logros Esperados' not in cab: continue
            if len(table.rows) < 2:
                continue
            anio_columnas = ajustar_columnas_anio(
                table, anios_periodo, header_row=1, first_year_col=5
            )
            fallback_ordinals = {}
            for row in table.rows:
                cs = row.cells
                if len(cs) < 6: continue
                cod = cs[0].text.strip()
                m = re.match(r'^(OEI\.\d{2}|AEI\.\d{2}\.\d{2})$', cod)
                if not m:
                    continue
                codigo = m.group()
                ordinal = indicador_ordinal_por_fila.get(id(row._tr))
                if ordinal is None:
                    ordinal = indicador_ordinal_por_clave.get(
                        (codigo, tuple(cell.text.strip() for cell in cs))
                    )
                if ordinal is None:
                    ordinal = fallback_ordinals.get(codigo, 0)
                    fallback_ordinals[codigo] = ordinal + 1
                meta = metas_por_indicador.get((codigo, ordinal))
                if meta is None:
                    continue
                esc(cs[3], fmt(get(meta, ['año_base', 'anio_base'])))
                esc(cs[4], fmt(get(meta, ['valor_base'])))
                for year, column in anio_columnas.items():
                    if column >= len(cs):
                        continue
                    esc(cs[column], fmt(get(meta, [f'meta_{year}'])))
                mok += 1
        print(f"  Metas: {mok}")
        doc.save(temp1)

        # PASO 3: Fichas EN ORDEN JERÁRQUICO
        print("\nPASO 3: Fichas (orden jerárquico)...")
        fichas_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'fichas_tecnicas.json')
        with open(fichas_path, 'r', encoding='utf-8') as f:
            fichas_db = json.load(f)
        fichas_index = indexar_fichas(fichas_db)

        doc_base = Document(temp1)
        doc_plt = Document(ficha_template_path)
        tbl_plt = doc_plt.tables[0]
        nota_ficha = extraer_nota_ficha(tbl_plt._element)
        ficha_anio_columnas = ajustar_columnas_anio(
            tbl_plt, anios_periodo, header_row=11, first_year_col=2
        )

        fichas_a_generar = construir_fichas_seleccionadas(
            codigos_ordenados, indices_oei, indices_aei, fichas_index
        )

        body = doc_base.element.body
        eliminar_parrafos_con_texto(body, 'Agregar fichas de indicadores')
        anexo_title = buscar_parrafo_con_texto(body, ANEXO_FICHAS_MARKER)
        if fichas_a_generar:
            if anexo_title is None:
                anexo_title = doc_base.add_paragraph('ANEXO A - 6: FICHA TÉCNICA DE INDICADORES')._p
            eliminar_parrafos_vacios_adyacentes(body, anexo_title)
            asegurar_salto_de_pagina_antes(anexo_title)
        elif anexo_title is not None:
            eliminar_parrafos_vacios_adyacentes(body, anexo_title)
            body.remove(anexo_title)

        fg = 0
        # USAR ORDEN JERÁRQUICO
        for codigo, ordinal in fichas_a_generar:
            ficha = seleccionar_ficha_base(fichas_index, codigo, ordinal, anios_periodo)
            fichas_anuales = obtener_fichas_indicador(fichas_index, codigo, ordinal)
            if ficha is None:
                print(f"  Aviso: {codigo} indicador {ordinal} no tiene ficha en JSON")
                continue

            # Copiar tabla
            nt = deepcopy(tbl_plt._element)
            nota = deepcopy(nota_ficha) if nota_ficha is not None else None
            if nota is not None:
                _set_xml_text(nota, obtener_nota_ficha(ficha) or _xml_text(nota_ficha))
            insertar_antes_de_sectpr(body, nt)
            if nota is not None:
                insertar_despues_de_elemento(body, nt, nota)
            tb = doc_base.tables[-1]
            if fg > 0 and tb.rows and tb.rows[0].cells and tb.rows[0].cells[0].paragraphs:
                tb.rows[0].cells[0].paragraphs[0].paragraph_format.page_break_before = True
            
            # Llenar datos
            def stxt(ri, ci, v):
                if ri < len(tb.rows) and ci < len(tb.rows[ri].cells):
                    cell = tb.rows[ri].cells[ci]
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = ''
                    if cell.paragraphs:
                        if cell.paragraphs[0].runs:
                            cell.paragraphs[0].runs[0].text = str(v)
                        else:
                            cell.paragraphs[0].add_run(str(v))
            
            stxt(1, 1, ficha.get('objetivo_accion', ''))
            stxt(2, 1, ficha.get('nombre_indicador', ''))
            stxt(3, 1, ficha.get('justificacion', ''))
            stxt(4, 1, ficha.get('responsables', ''))
            stxt(5, 1, ficha.get('limitaciones', ''))
            stxt(6, 1, ficha.get('metodo_calculo', ''))
            stxt(7, 1, ficha.get('sentido_esperado', ''))
            stxt(8, 1, ficha.get('proceso_recoleccion', ''))
            stxt(9, 1, ficha.get('fuente_datos', ''))
            stxt(11, 1, ficha.get('anio_base', ficha.get('año_base', '')))
            stxt(12, 1, ficha.get('valor_relativo', ''))
            stxt(13, 1, valor_absoluto_ficha(ficha))
            for column in ficha_anio_columnas.values():
                stxt(12, column, '')
                stxt(13, column, '')
            for year, column in ficha_anio_columnas.items():
                anual = fichas_anuales.get(year)
                if anual is None:
                    continue
                stxt(12, column, anual.get('valor_relativo', ''))
                stxt(13, column, valor_absoluto_ficha(anual))
            
            fg += 1
            print(f"  {fg}. {codigo} indicador {ordinal}")

        print(f"\nTotal fichas generadas: {fg}")

        output = output_path or f"PEI_{sanitizar_nombre_municipio(d['nombre_municipio'])}.docx"
        limpiar_parrafos_body_vacios_finales(doc_base.element.body)
        doc_base.save(output)

        print(f"\nGENERADO: {output}")
        print("="*70)
        return str(output)

def iniciar_servidor():
    httpd = HTTPServer(('', 8000), PEIHandler)
    print("\n" + "="*70)
    print("  GENERADOR PEI - ORDEN JERÁRQUICO")
    print("  http://localhost:8000")
    print("="*70 + "\n")
    try:
        httpd.serve_forever()
    except KeyboardInterrupt:
        print("\nDetenido.")


if __name__ == '__main__':
    iniciar_servidor()
